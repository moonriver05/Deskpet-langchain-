import os
import sys
import json
import time
import requests
import datetime
import random
import base64
import re
import pymysql
import math
import concurrent.futures
import threading
import uuid
import ctypes
import ctypes.wintypes
import tempfile
import PyQt5
from openai import OpenAI
from langchain_openai import ChatOpenAI
from langchain_core.messages import HumanMessage, SystemMessage
from zhdate import ZhDate
from pet_core.config import CONFIG_SCHEMA, REQUIRED_KEYS, app_config
from pet_core import window_chrome
from pet_core.persona import (
    ALICE_PROACTIVE_PERSONA,
    ALICE_RESPONSE_STYLE,
    ALICE_RUNTIME_CAPABILITIES,
    DEFAULT_CATEGORY_DESCRIPTIONS,
    build_current_response_card,
    format_capability_registry_for_prompt,
)
from pet_core.pending_intent import (
    PendingIntentStore,
    build_pending_intent_from_reply,
)
from pet_core.proactive import (
    clean_proactive_message,
    format_proactive_context_for_prompt,
    is_repetitive_proactive_message,
)
from pet_core.auto_labeler import (
    schedule_auto_label_batch,
    schedule_auto_relabel_event,
)
from pet_core.learning_logger import (
    build_interaction_event,
    log_feedback_event,
    log_implicit_state_observation,
    log_interaction_event,
)
from pet_core.knowledge_router import should_search_knowledge_base
from pet_core.recommender import (
    format_recommendation_for_prompt,
    recommendation_runtime,
)
from pet_core.rss_content import (
    format_external_content_for_prompt,
    format_external_content_recommendation_message,
    rss_content_runtime,
)
from pet_core.rsshub_service import (
    configure_local_rsshub_base,
    start_local_rsshub_background,
)
from pet_core.strategy_predictor import strategy_predictor_runtime
from pet_core.system_state import collect_system_state
from pet_core.timer_parser import (
    format_focus_duration,
)
from pet_features.todo_system import (
    TodoToolRouterThread,
    TodoWindow,
    has_explicit_todo_write_intent,
    todo_store,
)
from pet_features.learning_label_window import LearningLabelWindow
from pet_features.rss_manager_window import RSSManagerWindow
from pet_memory.memory_system import (
    DB_CONFIG,
    DB_NAME,
    MEM_IMP_CAP,
    MEM_PROMOTE_TO_LONG_TERM_SCORE,
    MEM_REPEATED_BUMP,
    _clean_profile_claim_text,
    _initial_importance_for_memory,
    _is_bad_profile_claim,
    _normalize_profile_category,
    _profile_bucket_for_memory,
    _profile_refiner_enabled,
    _refresh_user_profile_from_claims,
    configure_memory_database,
    conversation_history,
    daily_decay_memory,
    delete_profile_evidence_for_source,
    delete_short_memory_from_chroma,
    get_db_name,
    get_user_profile_prompt_context,
    init_db,
    knowledge_base,
    memory_runtime,
    promote_memory_to_long_term,
    refresh_user_profile_from_long_term,
    schedule_chroma_sync_repair,
    schedule_profile_refine,
    schedule_refine_unprocessed_long_term_memories,
    soul_state,
    sync_short_memory_to_chroma,
)
from pet_services.chroma_mcp import (
    CHROMA_COLLECTION_MEM,
    chrom_distance_to_sim,
    chroma_add_documents_sync,
    chroma_delete_documents_sync,
    chroma_query_documents_sync,
    configure_chroma,
)
try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

try:
    import docx
except ImportError:
    docx = None

try:
    import jieba
except ImportError:
    jieba = None

try:
    from qcloud_cos import CosConfig
    from qcloud_cos import CosS3Client
except ImportError:
    pass


# ====================  (AppConfig) ====================
def apply_config_to_globals():
    """把 app_config 的值同步到老的模块级全局变量（DB_CONFIG / COS_CONFIG / ark_api_key 等）。
    设置窗口保存后会再调一次，使新值即时生效（不用重启）。
    """
    global ark_api_key

    mysql_cfg = app_config.get_section("mysql")
    configure_memory_database(mysql_cfg)
    globals()["DB_NAME"] = get_db_name()

    cos_cfg = app_config.get_section("cos")
    COS_CONFIG["secret_id"]  = cos_cfg.get("secret_id")  or ""
    COS_CONFIG["secret_key"] = cos_cfg.get("secret_key") or ""
    COS_CONFIG["region"]     = cos_cfg.get("region")     or "ap-guangzhou"
    COS_CONFIG["bucket"]     = cos_cfg.get("bucket")     or ""
    COS_CONFIG["base_url"]   = cos_cfg.get("base_url")   or ""

    chroma_cfg = app_config.get_section("chroma")
    configure_chroma(chroma_cfg.get("container_name"))

    ark_api_key = app_config.get("ark.api_key", "") or ""

    tts_cfg = app_config.get_section("tts")
    configure_tts(tts_cfg, base_dir=os.path.dirname(os.path.abspath(__file__)))


# 提前声明 ark_api_key，让 apply_config_to_globals 第一次调用时能 import-time 赋值
ark_api_key = ""



# ==================== 腾讯云 COS 图床配置 ====================
# 真实值来自 AppConfig。留空时 EmotionCOSManager 会自己降级（不发表情包），
# SettingsWindow 保存后通过 apply_config_to_globals() 原地刷新。
COS_CONFIG = {
    'secret_id':  app_config.get("cos.secret_id", "")  or "",
    'secret_key': app_config.get("cos.secret_key", "") or "",
    'region':     app_config.get("cos.region", "ap-guangzhou") or "ap-guangzhou",
    'bucket':     app_config.get("cos.bucket", "")    or "",
    'base_url':   app_config.get("cos.base_url", "")  or "",
}

class EmotionCOSManager:
    def __init__(self):
        self.client = None
        self.bucket = COS_CONFIG['bucket']
        self.base_url = COS_CONFIG['base_url']
        if COS_CONFIG['secret_id'] and COS_CONFIG['secret_key'] and self.bucket:
            try:
                config = CosConfig(Region=COS_CONFIG['region'], SecretId=COS_CONFIG['secret_id'], SecretKey=COS_CONFIG['secret_key'])
                self.client = CosS3Client(config)
            except Exception as e:
                print(f"COS 初始化失败: {e}")

    def sync_local_memes(self, local_dir):
        """同步本地 memes 文件夹到 COS 图床"""
        if not self.client:
            return False, "COS 未初始化"
            
        if not os.path.exists(local_dir):
            os.makedirs(local_dir, exist_ok=True)
            return True, f"本地文件夹不存在，已创建：{local_dir}\n请将表情包放入对应的英文情感文件夹中！"

        success_count = 0
        skip_count = 0
        error_count = 0
        
        try:
            # 获取云端已有的所有文件列表，避免重复上传
            remote_files = set()
            marker = ""
            while True:
                resp = self.client.list_objects(Bucket=self.bucket, Marker=marker, MaxKeys=1000)
                if 'Contents' in resp:
                    for item in resp['Contents']:
                        remote_files.add(item['Key'])
                if resp.get('IsTruncated') == 'true':
                    marker = resp['NextMarker']
                else:
                    break
            
            # 遍历本地文件夹
            for emotion_dir in os.listdir(local_dir):
                emotion_path = os.path.join(local_dir, emotion_dir)
                if os.path.isdir(emotion_path):
                    for filename in os.listdir(emotion_path):
                        if filename.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.webp')):
                            local_file_path = os.path.join(emotion_path, filename)
                            # 构造在 COS 中的路径（例如 "angry/1.gif"）
                            cos_key = f"{emotion_dir}/{filename}"
                            
                            if cos_key in remote_files:
                                skip_count += 1
                                continue
                                
                            try:
                                self.client.upload_file(
                                    Bucket=self.bucket,
                                    LocalFilePath=local_file_path,
                                    Key=cos_key
                                )
                                success_count += 1
                            except Exception as upload_e:
                                print(f"上传 {cos_key} 失败: {upload_e}")
                                error_count += 1
                                
            return True, f"同步完成！\n成功上传: {success_count} 张\n跳过已存在: {skip_count} 张\n失败: {error_count} 张"
        except Exception as e:
            return False, f"同步过程发生错误: {str(e)}"

    def get_random_emotion_image(self, emotion):
        if not self.client:
            return None
        try:
            # 假设图床中表情包存放在以情感命名的文件夹下，例如 "开心/"
            response = self.client.list_objects(
                Bucket=self.bucket,
                Prefix=f"{emotion}/"
            )
            if 'Contents' in response:
                files = [item['Key'] for item in response['Contents'] if item['Key'].lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.webp'))]
                if files:
                    selected = random.choice(files)
                    return f"{self.base_url.rstrip('/')}/{selected}"
        except Exception as e:
            print(f"COS 获取表情包失败: {e}")
        return None

cos_manager = EmotionCOSManager()

# 关键：手动指定 Qt 插件目录
plugin_path = os.path.join(os.path.dirname(PyQt5.__file__), "Qt5", "plugins")
os.environ["QT_QPA_PLATFORM_PLUGIN_PATH"] = plugin_path

from PyQt5.QtWidgets import (
    QApplication, QWidget, QLabel, QVBoxLayout, QHBoxLayout,
    QPushButton, QLineEdit, QListWidget, QListWidgetItem,
    QCheckBox, QMessageBox, QMenu, QScrollArea, QDesktopWidget,
    QFileDialog, QFrame, QComboBox,
    QSizePolicy, QStackedWidget, QDialog, QFormLayout, QSpinBox,
    QPlainTextEdit, QTextEdit, QSplitter, QTabWidget, QTableWidget,
    QTableWidgetItem, QHeaderView, QAbstractItemView, QProgressBar
)
from PyQt5.QtCore import (
    Qt, QPoint, QTimer, QSize, QThread, pyqtSignal, QObject,
    QPropertyAnimation, QEasingCurve, QRect, QUrl, QLockFile,
    qInstallMessageHandler
)
from PyQt5.QtGui import (
    QPixmap, QPainter, QFont, QColor, QPolygon, QMovie, QIcon,
    QLinearGradient, QBrush, QPen, QDesktopServices
)
from pet_services.tts_service import (
    TTSSynthThread,
    cleanup_tts_artifacts,
    configure_tts,
    play_tts_file,
    reset_tts_client_state,
)


def apply_dark_window_chrome(widget):
    """设置无边框窗口和透明背景。

    注意：不要在顶层透明窗口上挂 QGraphicsDropShadowEffect。Windows 的
    layered window 会把阴影区域计入 dirty rect，若超出窗口边界就会反复
    打印 UpdateLayeredWindowIndirect failed。
    """
    widget.setWindowFlags(widget.windowFlags() | Qt.Window | Qt.FramelessWindowHint)
    widget.setAttribute(Qt.WA_TranslucentBackground)


def install_qt_message_filter():
    """Silence noisy Windows layered-window warnings without hiding real errors."""
    def _handler(mode, context, message):
        text = str(message or "")
        if "UpdateLayeredWindowIndirect failed" in text:
            return
        try:
            sys.stderr.write(text + "\n")
        except Exception:
            pass

    qInstallMessageHandler(_handler)


def acquire_single_instance_lock():
    """Prevent two desktop-pet processes from fighting over DB/RSS/Chroma/logs."""
    lock_dir = os.path.join(tempfile.gettempdir(), "YuzuDeskpet")
    os.makedirs(lock_dir, exist_ok=True)
    lock = QLockFile(os.path.join(lock_dir, "deskpet.lock"))
    lock.setStaleLockTime(30 * 1000)
    if not lock.tryLock(100):
        QMessageBox.information(
            None,
            "有珠已经在运行",
            "已经有一个桌宠进程在运行了。\n"
            "如果桌宠卡住了，请先关闭旧窗口或结束旧的 python pet.py 进程，再重新启动。",
        )
        return None
    return lock


def _is_interactive_widget(widget):
    while widget is not None:
        if isinstance(widget, (
            QPushButton,
            QLineEdit,
            QTextEdit,
            QPlainTextEdit,
            QComboBox,
            QSpinBox,
            QListWidget,
            QTableWidget,
            QScrollArea,
        )):
            return True
        widget = widget.parentWidget()
    return False


def _begin_window_drag(window, event, drag_widget=None, max_y=None):
    if event.button() != Qt.LeftButton:
        return False
    if drag_widget is not None:
        top_left = drag_widget.mapTo(window, QPoint(0, 0))
        if not drag_widget.rect().translated(top_left).contains(event.pos()):
            return False
    elif max_y is not None and event.pos().y() > max_y:
        return False
    if _is_interactive_widget(window.childAt(event.pos())):
        return False
    window._dragging = True
    window._drag_offset = event.globalPos() - window.frameGeometry().topLeft()
    event.accept()
    return True


def _continue_window_drag(window, event):
    if getattr(window, "_dragging", False) and getattr(window, "_drag_offset", None) is not None:
        if event.buttons() == Qt.LeftButton:
            window.move(event.globalPos() - window._drag_offset)
            event.accept()
            return True
    return False


def _end_window_drag(window):
    window._dragging = False
    window._drag_offset = None


# ark_api_key 已经在文件顶部声明，apply_config_to_globals() 会从 AppConfig 注入。
# 这里再 sync 一次，覆盖掉 import 时还没载入 AppConfig 的占位值（"" → 真实 key）。
ark_api_key = app_config.get("ark.api_key", "") or ""


# ==================== 设置窗口 (SettingsWindow) ====================
# 左侧分类列表 + 右侧表单堆叠，所有字段由 CONFIG_SCHEMA 驱动。
# 保存时把 QLineEdit / QSpinBox 里的当前值写回 app_config，再 .save() 落盘，
# 最后调一次 apply_config_to_globals() 让新值即时生效。
class SettingsWindow(QDialog):
    def __init__(self, parent=None, highlight_section=None):
        super().__init__(parent)
        self.setWindowTitle("桌宠设置")
        apply_dark_window_chrome(self)
        self.resize(820, 600)
        window_chrome.setup_resizable_frameless_window(self, minimum_size=(640, 460))
        self.maximize_btn = None
        self._dragging = False
        self._drag_offset = None
        self.setStyleSheet("""
            QDialog {
                background: transparent;
                color: #E0E0E0;
                font-family: 'Segoe UI', 'Microsoft YaHei', sans-serif;
            }
            QFrame#SettingsContainer {
                background: #1A1525;
                border: 1px solid #3D2E55;
                border-radius: 20px;
            }
            QFrame#SettingsTitleBar {
                background: transparent;
                border: none;
                border-bottom: 1px solid #2A203B;
            }
            QLabel#windowTitle {
                color: #B886F8;
                font-size: 14px;
                font-weight: 500;
                background: transparent;
            }
            QWidget { background: transparent; color: #E0E0E0; }
            QListWidget {
                background: #15151D;
                color: #A0A0B0;
                border: none;
                border-right: 1px solid #2A2A3C;
                padding: 10px;
                font-size: 14px;
                outline: 0;
            }
            QListWidget::item {
                padding: 12px 14px;
                border-radius: 16px;
                margin: 4px 0;
            }
            QListWidget::item:hover {
                background: #252535;
                color: #FFFFFF;
            }
            QListWidget::item:selected {
                background: rgba(169, 123, 255, 0.15);
                color: #A97BFF;
                border: 1px solid #A97BFF;
                font-weight: bold;
            }
            QLineEdit, QSpinBox, QPlainTextEdit {
                background: #252535;
                color: #FFFFFF;
                border: 1px solid #3D3D52;
                border-radius: 16px;
                padding: 8px 12px;
                font-size: 14px;
                selection-background-color: #A97BFF;
            }
            QLineEdit:focus, QSpinBox:focus, QPlainTextEdit:focus {
                border: 1px solid #A97BFF;
                background: #2A2A3C;
            }
            QLabel#desc { color: #A0A0B0; font-size: 13px; }
            QLabel#hint { color: #6D6D8A; font-size: 12px; font-style: italic; }
            QLabel#title { font-size: 18px; font-weight: bold; color: #A97BFF; }
            QPushButton {
                background: #252535;
                color: #E0E0E0;
                border: 1px solid #3D3D52;
                border-radius: 16px;
                padding: 8px 18px;
                font-weight: bold;
                font-size: 14px;
            }
            QPushButton:hover {
                background: #3D2E55;
                border-color: #A97BFF;
                color: #A97BFF;
            }
            QPushButton[role="ghost"] {
                background: transparent;
                color: #A0A0B0;
                border: 1px solid #3D3D52;
            }
            QPushButton[role="ghost"]:hover {
                background: #2A2A3C;
                color: #FFFFFF;
                border-color: #A97BFF;
            }
            QPushButton[role="title"] {
                background: transparent;
                color: #B8ADC9;
                border: none;
                border-radius: 12px;
                padding: 0;
                font-size: 13px;
                font-weight: bold;
            }
            QPushButton[role="title"]:hover {
                background: #2A203B;
                color: #EAE5F2;
            }
            QPushButton[role="eye"] {
                background: transparent;
                color: #A97BFF;
                border: none;
                padding: 0px; font-size: 14px;
            }
            QPushButton[role="eye"]:hover { color: #FFFFFF; }
            QPushButton[role="picker"] {
                background: #15151D;
                color: #E0E0E0;
                border: 1px solid #3D2E55;
                padding: 5px 11px;
                font-weight: normal;
            }
            QPushButton[role="picker"]:hover {
                background: #2A203B;
                border-color: #B886F8;
            }
            QScrollArea { background: transparent; border: none; }
            QScrollBar:vertical {
                background: #1A1525;
                width: 8px;
                margin: 0;
            }
            QScrollBar::handle:vertical {
                background: #3D2E55;
                border-radius: 16px;
                min-height: 24px;
            }
            QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical {
                height: 0;
                border: none;
            }
        """)

        # field_key (str like "mysql.password") → QWidget（输入控件）
        self._inputs = {}

        shell = QVBoxLayout(self)
        shell.setContentsMargins(20, 20, 20, 20)
        shell.setSpacing(0)
        self.container = QFrame(self)
        self.container.setObjectName("SettingsContainer")
        shell.addWidget(self.container)

        outer = QVBoxLayout(self.container)
        outer.setContentsMargins(0, 0, 0, 0)
        outer.setSpacing(0)

        self.title_bar = QFrame()
        self.title_bar.setObjectName("SettingsTitleBar")
        self.title_bar.setFixedHeight(40)
        title_layout = QHBoxLayout(self.title_bar)
        title_layout.setContentsMargins(16, 0, 10, 0)
        title_layout.setSpacing(8)
        title = QLabel("桌宠设置")
        title.setObjectName("windowTitle")
        title_layout.addWidget(title)
        title_layout.addStretch()

        min_btn = QPushButton("—")
        min_btn.setFixedSize(26, 24)
        min_btn.setToolTip("最小化")
        min_btn.setProperty("role", "title")
        min_btn.clicked.connect(self.showMinimized)
        title_layout.addWidget(min_btn)

        self.maximize_btn = QPushButton("□")
        self.maximize_btn.setFixedSize(26, 24)
        self.maximize_btn.setToolTip("最大化/还原")
        self.maximize_btn.setProperty("role", "title")
        self.maximize_btn.clicked.connect(lambda: window_chrome.toggle_maximize_restore(self))
        title_layout.addWidget(self.maximize_btn)

        close_btn = QPushButton("✕")
        close_btn.setFixedSize(26, 24)
        close_btn.setToolTip("关闭")
        close_btn.setProperty("role", "title")
        close_btn.clicked.connect(self.reject)
        title_layout.addWidget(close_btn)
        outer.addWidget(self.title_bar)

        body = QHBoxLayout()
        body.setContentsMargins(0, 0, 0, 0)
        body.setSpacing(0)

        # 左侧分类列表
        self.cat_list = QListWidget()
        self.cat_list.setFixedWidth(180)
        for section in CONFIG_SCHEMA:
            item = QListWidgetItem(section['title'])
            self.cat_list.addItem(item)
        self.cat_list.currentRowChanged.connect(self._on_cat_changed)

        # 右侧堆叠
        self.stack = QStackedWidget()
        self.stack.setStyleSheet("background: #1A1525;")
        for section in CONFIG_SCHEMA:
            page = self._build_section_page(section)
            self.stack.addWidget(page)

        body.addWidget(self.cat_list)
        body.addWidget(self.stack, 1)
        outer.addLayout(body, 1)

        # 底部按钮条
        footer = QHBoxLayout()
        footer.setContentsMargins(15, 10, 15, 15)
        self.status_label = QLabel("")
        self.status_label.setStyleSheet("color: #A08CD2; font-size: 12px;")
        footer.addWidget(self.status_label, 1)

        btn_reset = QPushButton("恢复默认")
        btn_reset.setProperty("role", "ghost")
        btn_reset.clicked.connect(self._on_reset_clicked)

        btn_cancel = QPushButton("取消")
        btn_cancel.setProperty("role", "ghost")
        btn_cancel.clicked.connect(self.reject)

        btn_save = QPushButton("保存")
        btn_save.clicked.connect(self._on_save_clicked)

        footer.addWidget(btn_reset)
        footer.addWidget(btn_cancel)
        footer.addWidget(btn_save)
        outer.addLayout(footer)

        # 定位默认选中
        target_row = 0
        if highlight_section:
            for i, sec in enumerate(CONFIG_SCHEMA):
                if sec["key"] == highlight_section:
                    target_row = i
                    break
        self.cat_list.setCurrentRow(target_row)

    def showEvent(self, event):
        super().showEvent(event)
        window_chrome.sync_maximize_button(self)

    def nativeEvent(self, eventType, message):
        handled = window_chrome.native_resize_event(self, eventType, message)
        if handled is not None:
            return handled
        return super().nativeEvent(eventType, message)

    def mousePressEvent(self, event):
        if window_chrome.begin_window_resize(self, event):
            return
        if window_chrome.begin_title_drag(self, event, getattr(self, "title_bar", None)):
            return
        super().mousePressEvent(event)

    def mouseMoveEvent(self, event):
        if window_chrome.continue_window_resize(self, event):
            return
        if window_chrome.continue_title_drag(self, event):
            return
        window_chrome.update_resize_cursor(self, event)
        super().mouseMoveEvent(event)

    def mouseReleaseEvent(self, event):
        window_chrome.end_window_resize(self)
        window_chrome.end_title_drag(self)
        super().mouseReleaseEvent(event)

    def mouseDoubleClickEvent(self, event):
        if window_chrome.title_double_click_maximize(self, event, getattr(self, "title_bar", None)):
            return
        super().mouseDoubleClickEvent(event)

    def leaveEvent(self, event):
        window_chrome.leave_resize_area(self, event)
        super().leaveEvent(event)

    def resizeEvent(self, event):
        window_chrome.sync_maximize_button(self)
        super().resizeEvent(event)

    # ---------- 构造每个分类的表单 ----------
    def _build_section_page(self, section):
        page = QWidget()
        v = QVBoxLayout(page)
        v.setContentsMargins(24, 20, 24, 20)
        v.setSpacing(14)

        title = QLabel(section['title'])
        title.setObjectName("title")
        v.addWidget(title)

        if section.get("desc"):
            desc = QLabel(section["desc"])
            desc.setObjectName("desc")
            desc.setWordWrap(True)
            v.addWidget(desc)

        form_holder = QScrollArea()
        form_holder.setWidgetResizable(True)
        form_holder.setFrameShape(QFrame.NoFrame)
        form_holder.setStyleSheet("background: transparent;")

        form_widget = QWidget()
        form_widget.setStyleSheet("background: transparent;")
        form = QFormLayout(form_widget)
        form.setLabelAlignment(Qt.AlignRight)
        form.setFormAlignment(Qt.AlignTop)
        form.setSpacing(10)

        for field in section["fields"]:
            full_key = f"{section['key']}.{field['key']}"
            current = app_config.get(full_key, field.get("default", ""))
            row_widget, input_widget = self._build_field_row(field, current)
            self._inputs[full_key] = input_widget

            label_text = field["label"]
            if field.get("required"):
                label_text += " *"
            label = QLabel(label_text)
            label.setStyleSheet("font-size: 13px; color: #D1C8E1;")
            form.addRow(label, row_widget)

        form_holder.setWidget(form_widget)
        v.addWidget(form_holder, 1)

        return page

    def _build_field_row(self, field, current_value):
        """根据 field schema 构造一行：可能是 QLineEdit / QSpinBox / 路径选择 / 密码（带眼睛）。
        返回 (容器 widget, 输入 widget)。"""
        ftype = field.get("type", "str")
        if ftype == "int":
            inp = QSpinBox()
            inp.setMaximum(999999)
            inp.setMinimum(0)
            try:
                inp.setValue(int(current_value))
            except (TypeError, ValueError):
                inp.setValue(int(field.get("default", 0) or 0))
            row = QWidget()
            row.setStyleSheet("background: transparent;")
            row_l = QVBoxLayout(row)
            row_l.setContentsMargins(0, 0, 0, 0)
            row_l.setSpacing(2)
            row_l.addWidget(inp)
            if field.get("hint"):
                hint = QLabel(field["hint"])
                hint.setObjectName("hint")
                hint.setWordWrap(True)
                row_l.addWidget(hint)
            return row, inp

        if ftype == "text":
            inp = QPlainTextEdit()
            inp.setPlainText(str(current_value) if current_value is not None else "")
            inp.setPlaceholderText(str(field.get("default", "")))
            inp.setMinimumHeight(96)
            row = QWidget()
            row.setStyleSheet("background: transparent;")
            row_l = QVBoxLayout(row)
            row_l.setContentsMargins(0, 0, 0, 0)
            row_l.setSpacing(2)
            row_l.addWidget(inp)
            if field.get("hint"):
                hint = QLabel(field["hint"])
                hint.setObjectName("hint")
                hint.setWordWrap(True)
                row_l.addWidget(hint)
            return row, inp

        inp = QLineEdit()
        inp.setText(str(current_value) if current_value is not None else "")
        if field.get("secret"):
            inp.setEchoMode(QLineEdit.Password)
            inp.setPlaceholderText("（已隐藏）")
        else:
            inp.setPlaceholderText(str(field.get("default", "")))

        row = QWidget()
        row.setStyleSheet("background: transparent;")
        row_l = QVBoxLayout(row)
        row_l.setContentsMargins(0, 0, 0, 0)
        row_l.setSpacing(2)

        # 第一行：输入框 + 可选的眼睛/选文件按钮
        line = QHBoxLayout()
        line.setContentsMargins(0, 0, 0, 0)
        line.setSpacing(6)
        line.addWidget(inp, 1)

        if field.get("secret"):
            eye = QPushButton("👁")
            eye.setProperty("role", "eye")
            eye.setFixedWidth(28)
            eye.setCheckable(True)
            def _toggle(checked, _inp=inp, _eye=eye):
                _inp.setEchoMode(QLineEdit.Normal if checked else QLineEdit.Password)
                _eye.setText("🙈" if checked else "👁")
            eye.toggled.connect(_toggle)
            line.addWidget(eye)

        if ftype == "path":
            pick = QPushButton("浏览…")
            pick.setProperty("role", "picker")
            def _pick(_=False, _inp=inp):
                p, _filt = QFileDialog.getOpenFileName(
                    self, "选择参考音频", "",
                    "音频 (*.wav *.mp3 *.flac *.ogg);;全部 (*.*)",
                )
                if p:
                    _inp.setText(p)
            pick.clicked.connect(_pick)
            line.addWidget(pick)

        row_l.addLayout(line)

        if field.get("hint"):
            hint = QLabel(field["hint"])
            hint.setObjectName("hint")
            hint.setWordWrap(True)
            row_l.addWidget(hint)

        return row, inp

    # ---------- 事件 ----------
    def _on_cat_changed(self, idx):
        if idx >= 0:
            self.stack.setCurrentIndex(idx)

    def _collect_inputs(self):
        """读取所有输入控件 → 返回扁平的 {dot_path: value} dict。"""
        out = {}
        for full_key, widget in self._inputs.items():
            if isinstance(widget, QSpinBox):
                out[full_key] = int(widget.value())
            elif isinstance(widget, QPlainTextEdit):
                out[full_key] = widget.toPlainText().strip()
            elif isinstance(widget, QLineEdit):
                out[full_key] = widget.text().strip()
            else:
                out[full_key] = widget.text() if hasattr(widget, "text") else ""
        return out

    def _validate(self, flat):
        """检查必填项 + 个别格式。返回缺失的中文提示列表。"""
        problems = []
        for req_key in REQUIRED_KEYS:
            if not str(flat.get(req_key, "")).strip():
                # 找到 schema 里的中文 label
                for sec in CONFIG_SCHEMA:
                    if not req_key.startswith(sec["key"] + "."):
                        continue
                    fk = req_key.split(".", 1)[1]
                    for f in sec["fields"]:
                        if f["key"] == fk:
                            problems.append(f"{sec['title']} / {f['label']}")
                            break
                    break
        return problems

    def _on_save_clicked(self):
        flat = self._collect_inputs()
        problems = self._validate(flat)
        if problems:
            self.status_label.setText("以下必填项还没填：" + "、".join(problems))
            QMessageBox.warning(
                self, "还差一点",
                "下面这些是桌宠跑起来最低需要的字段，麻烦填一下：\n\n  · " +
                "\n  · ".join(problems),
            )
            return

        # 写回 AppConfig 并落盘
        for full_key, val in flat.items():
            app_config.set(full_key, val)
        ok = app_config.save()

        # 同步到老的全局变量
        try:
            apply_config_to_globals()
        except Exception as e:
            print(f"[Settings] apply_config_to_globals 失败：{e}")

        # 让 COS / TTS 这种"已经实例化过的"东西也尽量刷新
        self._reload_dependent_singletons()

        if ok:
            self.status_label.setText("")
            QMessageBox.information(
                self, "已保存",
                "已保存到 pet_config.json。\n"
                "大部分配置即时生效；如果是首次填 MySQL 密码或 Chroma 容器名，"
                "建议重启一次桌宠以确保所有连接刷新。",
            )
            self.accept()
        else:
            self.status_label.setText("保存失败（写文件出错）")

    def _reload_dependent_singletons(self):
        """COS / TTS 客户端有内部状态，刷一下让新配置生效。"""
        try:
            # cos_manager 重新初始化（用最新的 COS_CONFIG）
            globals()["cos_manager"] = EmotionCOSManager()
        except Exception as e:
            print(f"[Settings] 重建 cos_manager 失败：{e}")
        try:
            reset_tts_client_state()
        except Exception as e:
            print(f"[Settings] 重置 TTS weights 标记失败：{e}")
        try:
            configure_local_rsshub_base()
            start_local_rsshub_background()
        except Exception as e:
            print(f"[Settings] 应用本地 RSSHub 配置失败：{e}")

    def _on_reset_clicked(self):
        confirm = QMessageBox.question(
            self, "确认", "把所有字段恢复到默认值（不会立即保存）？",
            QMessageBox.Yes | QMessageBox.No,
        )
        if confirm != QMessageBox.Yes:
            return
        for section in CONFIG_SCHEMA:
            for field in section["fields"]:
                full_key = f"{section['key']}.{field['key']}"
                widget = self._inputs.get(full_key)
                if widget is None:
                    continue
                default = field.get("default", "")
                if isinstance(widget, QSpinBox):
                    try:
                        widget.setValue(int(default))
                    except (TypeError, ValueError):
                        widget.setValue(0)
                elif isinstance(widget, QPlainTextEdit):
                    widget.setPlainText(str(default) if default is not None else "")
                elif isinstance(widget, QLineEdit):
                    widget.setText(str(default) if default is not None else "")


def open_settings_dialog(parent=None, highlight_section=None):
    """供菜单/启动检查共用的入口。"""
    dlg = SettingsWindow(parent=parent, highlight_section=highlight_section)
    return dlg.exec_()


def ensure_required_config_or_prompt(parent=None):
    """启动时调用：若必填项缺失，弹设置窗口让用户填。
       用户填完且不再缺失 → 返回 True；用户取消且仍缺失 → 返回 False。"""
    while True:
        missing = app_config.missing_required()
        if not missing:
            return True
        labels = []
        for k in missing:
            for sec in CONFIG_SCHEMA:
                if not k.startswith(sec["key"] + "."):
                    continue
                fk = k.split(".", 1)[1]
                for f in sec["fields"]:
                    if f["key"] == fk:
                        labels.append(f"{sec['title']} / {f['label']}")
                        break
                break
        reply = QMessageBox.question(
            parent, "首次启动配置",
            "桌宠还没拿到这几项必要配置，没有它们没法工作：\n\n  · "
            + "\n  · ".join(labels)
            + "\n\n现在去填一下吗？（取消则跳过，桌宠仍会启动但相应功能不可用）",
            QMessageBox.Yes | QMessageBox.No,
        )
        if reply != QMessageBox.Yes:
            return False
        # 跳到第一个缺失字段所在分类
        first_section = missing[0].split(".", 1)[0]
        open_settings_dialog(parent=parent, highlight_section=first_section)
        # 循环回去再检查一次



# ==================== 桌宠主体 ====================
class MemoryManagerWindow(QWidget):
    """Audit and edit short-term memory, long-term memory, and user profile."""

    def __init__(self, pet=None):
        super().__init__()
        self.pet = pet
        self.current_rows = []
        self._dragging = False
        self._drag_offset = None
        self.setWindowTitle("记忆管理")
        apply_dark_window_chrome(self)
        self.resize(980, 680)
        window_chrome.setup_resizable_frameless_window(self, minimum_size=(760, 520))
        self.maximize_btn = None
        
        # 同样包裹 QFrame 以实现发光阴影与圆角
        self.container = QFrame(self)
        self.container.setObjectName("MainContainer")
        
        main_layout = QVBoxLayout(self)
        main_layout.setContentsMargins(20, 20, 20, 20)
        main_layout.addWidget(self.container)
        
        self.init_ui()
        self.refresh()

    def showEvent(self, event):
        super().showEvent(event)
        
        # 窗口弹出动画
        self.setWindowOpacity(0.0)
        self.anim = QPropertyAnimation(self, b"windowOpacity")
        self.anim.setDuration(250)
        self.anim.setStartValue(0.0)
        self.anim.setEndValue(1.0)
        self.anim.start()
        window_chrome.sync_maximize_button(self)

    def init_ui(self):
        self.setStyleSheet("""
            QFrame#MainContainer {
                background: #1A1525;
                border-radius: 20px;
                border: 1px solid #3D2E55;
            }
            QFrame#MemoryTitleBar {
                background: transparent;
                border: none;
                border-bottom: 1px solid #2A203B;
            }
            QWidget { background: transparent; color: #EAE5F2; font-family: 'Segoe UI', 'Microsoft YaHei', sans-serif; }
            QLabel { background: transparent; color: #B8ADC9; }
            QLineEdit, QPlainTextEdit, QComboBox {
                background: #231B32;
                color: #EAE5F2;
                border: 1px solid #4E3C6B;
                border-radius: 16px;
                padding: 7px 9px;
                font-size: 13px;
                selection-background-color: #4E3C6B;
            }
            QLineEdit:focus, QPlainTextEdit:focus, QComboBox:focus {
                border-color: #B886F8;
                background: #1A1525;
            }
            QComboBox::drop-down { width: 18px; border: none; }
            QPushButton {
                background: #2A203B;
                color: #EAE5F2;
                border: 1px solid #4E3C6B;
                border-radius: 16px;
                padding: 7px 13px;
                font-weight: 500;
            }
            QPushButton:hover {
                background: #3D2E55;
                border-color: #B886F8;
            }
            QPushButton[role="danger"] {
                background: #4E3C6B;
                border-color: #B886F8;
                color: #EAE5F2;
            }
            QPushButton[role="danger"]:hover { background: #634D85; }
            QPushButton[role="ghost"] {
                background: transparent;
                color: #D1C8E1;
                border: 1px solid #3D2E55;
            }
            QTableWidget {
                background: #1A1525;
                alternate-background-color: #231B32;
                color: #EAE5F2;
                border: 1px solid #3D2E55;
                border-radius: 16px;
                gridline-color: #231B32;
                selection-background-color: rgba(159, 183, 204, 58);
                selection-color: #FFFFFF;
            }
            QHeaderView::section {
                background: #3D2E55;
                color: #EAE5F2;
                border: none;
                border-right: 1px solid #4E3C6B;
                padding: 7px;
                font-weight: 600;
            }
            QScrollBar:vertical {
                background: #1A1525;
                width: 8px;
                margin: 0;
            }
            QScrollBar::handle:vertical {
                background: #3D2E55;
                border-radius: 16px;
                min-height: 24px;
            }
            QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical {
                height: 0;
                border: none;
            }
        """)

        root = QVBoxLayout(self.container)
        root.setContentsMargins(14, 14, 14, 14)
        root.setSpacing(10)

        self.header_bar = QFrame()
        self.header_bar.setObjectName("MemoryTitleBar")
        self.header_bar.setFixedHeight(42)
        header = QHBoxLayout(self.header_bar)
        header.setContentsMargins(2, 0, 2, 0)
        header.setSpacing(8)
        title = QLabel("记忆管理")
        title.setFont(QFont("Microsoft YaHei", 17, QFont.Light))
        title.setStyleSheet("color: #B886F8; letter-spacing: 0px;")
        header.addWidget(title)
        header.addStretch()
        refresh_profile_btn = QPushButton("刷新画像")
        refresh_profile_btn.setToolTip("重新从长期记忆聚合用户画像")
        refresh_profile_btn.clicked.connect(self.refresh_profile)
        header.addWidget(refresh_profile_btn)
        refine_pending_btn = QPushButton("精炼未处理")
        refine_pending_btn.setToolTip("最多处理 10 条还没进入 Profile Refiner 的长期记忆；会调用画像精炼模型")
        refine_pending_btn.clicked.connect(self.refine_pending_memories)
        header.addWidget(refine_pending_btn)

        window_btn_style = """
            QPushButton {
                background: transparent;
                color: #B8ADC9;
                border: none;
                border-radius: 12px;
                padding: 0;
                font-size: 13px;
                font-weight: bold;
            }
            QPushButton:hover {
                background: #2A203B;
                color: #EAE5F2;
            }
        """
        min_btn = QPushButton("-")
        min_btn.setFixedSize(26, 24)
        min_btn.setToolTip("最小化")
        min_btn.setStyleSheet(window_btn_style)
        min_btn.clicked.connect(self.showMinimized)
        header.addWidget(min_btn)

        self.maximize_btn = QPushButton("□")
        self.maximize_btn.setFixedSize(26, 24)
        self.maximize_btn.setToolTip("最大化/还原")
        self.maximize_btn.setStyleSheet(window_btn_style)
        self.maximize_btn.clicked.connect(lambda: window_chrome.toggle_maximize_restore(self))
        header.addWidget(self.maximize_btn)

        close_title_btn = QPushButton("×")
        close_title_btn.setFixedSize(26, 24)
        close_title_btn.setToolTip("关闭")
        close_title_btn.setStyleSheet(window_btn_style)
        close_title_btn.clicked.connect(self.close)
        header.addWidget(close_title_btn)
        root.addWidget(self.header_bar)

        filters = QHBoxLayout()
        self.mode_combo = QComboBox()
        self.mode_combo.addItem("短期记忆", "short")
        self.mode_combo.addItem("长期记忆", "long")
        self.mode_combo.addItem("用户画像", "profile")
        self.mode_combo.currentIndexChanged.connect(self.refresh)
        filters.addWidget(self.mode_combo)

        self.search_input = QLineEdit()
        self.search_input.setPlaceholderText("搜索内容、关键词、分类...")
        self.search_input.returnPressed.connect(self.refresh)
        filters.addWidget(self.search_input, 1)

        refresh_btn = QPushButton("搜索/刷新")
        refresh_btn.clicked.connect(self.refresh)
        filters.addWidget(refresh_btn)
        root.addLayout(filters)

        splitter = QSplitter(Qt.Vertical)
        root.addWidget(splitter, 1)

        self.table = QTableWidget()
        self.table.setAlternatingRowColors(True)
        self.table.setShowGrid(False)
        self.table.setSelectionBehavior(QAbstractItemView.SelectRows)
        self.table.setSelectionMode(QAbstractItemView.SingleSelection)
        self.table.setEditTriggers(QAbstractItemView.NoEditTriggers)
        self.table.verticalHeader().setVisible(False)
        self.table.horizontalHeader().setStretchLastSection(True)
        self.table.itemSelectionChanged.connect(self.on_selection_changed)
        splitter.addWidget(self.table)

        editor = QWidget()
        editor_layout = QVBoxLayout(editor)
        editor_layout.setContentsMargins(0, 8, 0, 0)
        editor_layout.setSpacing(8)

        meta_row = QHBoxLayout()
        self.selected_label = QLabel("未选择")
        self.selected_label.setStyleSheet("color: #B8ADC9;")
        meta_row.addWidget(self.selected_label)
        meta_row.addStretch()
        meta_row.addWidget(QLabel("分类/字段"))
        self.category_input = QLineEdit()
        self.category_input.setPlaceholderText("长期记忆分类或画像字段")
        self.category_input.setMaximumWidth(220)
        meta_row.addWidget(self.category_input)
        meta_row.addWidget(QLabel("关键词"))
        self.keywords_input = QLineEdit()
        self.keywords_input.setPlaceholderText("短期/长期关键词")
        self.keywords_input.setMaximumWidth(240)
        meta_row.addWidget(self.keywords_input)
        editor_layout.addLayout(meta_row)

        self.detail_edit = QPlainTextEdit()
        self.detail_edit.setPlaceholderText("选中一条记忆后可以在这里编辑正文/画像内容。")
        editor_layout.addWidget(self.detail_edit, 1)

        buttons = QHBoxLayout()
        self.clear_btn = QPushButton("清空表单")
        self.clear_btn.setProperty("role", "ghost")
        self.clear_btn.clicked.connect(self.clear_form)
        buttons.addWidget(self.clear_btn)

        self.add_btn = QPushButton("添加新项")
        self.add_btn.clicked.connect(self.add_new)
        buttons.addWidget(self.add_btn)

        self.save_btn = QPushButton("保存修改")
        self.save_btn.clicked.connect(self.save_selected)
        buttons.addWidget(self.save_btn)

        self.promote_btn = QPushButton("迁移为长期记忆")
        self.promote_btn.setToolTip("只对短期记忆有效：写入长期记忆后从短期表和 Chroma 删除")
        self.promote_btn.clicked.connect(self.promote_selected)
        buttons.addWidget(self.promote_btn)

        self.delete_btn = QPushButton("删除选中")
        self.delete_btn.setProperty("role", "danger")
        self.delete_btn.clicked.connect(self.delete_selected)
        buttons.addWidget(self.delete_btn)

        buttons.addStretch()
        close_btn = QPushButton("关闭")
        close_btn.setProperty("role", "ghost")
        close_btn.clicked.connect(self.close)
        buttons.addWidget(close_btn)
        editor_layout.addLayout(buttons)

        splitter.addWidget(editor)
        splitter.setSizes([420, 220])

    def _connect_db(self):
        return pymysql.connect(
            host=DB_CONFIG['host'], user=DB_CONFIG['user'],
            password=DB_CONFIG['password'], database=DB_NAME,
            charset=DB_CONFIG['charset'], cursorclass=pymysql.cursors.DictCursor
        )

    def mousePressEvent(self, event):
        if window_chrome.begin_window_resize(self, event):
            return
        if window_chrome.begin_title_drag(self, event, getattr(self, "header_bar", None)):
            return
        super().mousePressEvent(event)

    def mouseMoveEvent(self, event):
        if window_chrome.continue_window_resize(self, event):
            return
        if window_chrome.continue_title_drag(self, event):
            return
        window_chrome.update_resize_cursor(self, event)
        super().mouseMoveEvent(event)

    def mouseReleaseEvent(self, event):
        window_chrome.end_window_resize(self)
        window_chrome.end_title_drag(self)
        super().mouseReleaseEvent(event)

    def nativeEvent(self, eventType, message):
        handled = window_chrome.native_resize_event(self, eventType, message)
        if handled is not None:
            return handled
        return super().nativeEvent(eventType, message)

    def mouseDoubleClickEvent(self, event):
        if window_chrome.title_double_click_maximize(self, event, getattr(self, "header_bar", None)):
            return
        super().mouseDoubleClickEvent(event)

    def leaveEvent(self, event):
        window_chrome.leave_resize_area(self, event)
        super().leaveEvent(event)

    def resizeEvent(self, event):
        window_chrome.sync_maximize_button(self)
        super().resizeEvent(event)

    def _mode(self):
        return self.mode_combo.currentData() or "short"

    def refresh(self):
        mode = self._mode()
        q = self.search_input.text().strip()
        self.current_rows = []
        conn = None
        try:
            conn = self._connect_db()
            with conn.cursor() as cursor:
                like = f"%{q}%"
                if mode == "short":
                    sql = (
                        "SELECT id, content, keywords, importance_score, access_count, "
                        "created_at, last_accessed_at FROM user_memory "
                    )
                    args = []
                    if q:
                        sql += "WHERE content LIKE %s OR keywords LIKE %s "
                        args.extend([like, like])
                    sql += "ORDER BY importance_score DESC, last_accessed_at DESC LIMIT 300"
                    cursor.execute(sql, args)
                elif mode == "long":
                    sql = (
                        "SELECT id, source_memory_id, content, keywords, category, importance_score, "
                        "promote_reason, promoted_at FROM long_term_memory "
                    )
                    args = []
                    if q:
                        sql += "WHERE content LIKE %s OR keywords LIKE %s OR category LIKE %s "
                        args.extend([like, like, like])
                    sql += "ORDER BY promoted_at DESC, importance_score DESC LIMIT 300"
                    cursor.execute(sql, args)
                else:
                    _refresh_user_profile_from_claims(cursor, clear_when_empty=True)
                    sql = (
                        "SELECT id, field_name, claim, confidence, evidence_count, updated_at "
                        "FROM profile_claim "
                    )
                    args = []
                    if q:
                        sql += "WHERE field_name LIKE %s OR claim LIKE %s "
                        args.extend([like, like])
                    sql += "ORDER BY confidence DESC, updated_at DESC"
                    cursor.execute(sql, args)
                self.current_rows = list(cursor.fetchall())
            conn.commit()
        except Exception as e:
            QMessageBox.warning(self, "读取失败", f"读取记忆失败：{e}")
            self.current_rows = []
        finally:
            try:
                if conn:
                    conn.close()
            except Exception:
                pass
        self.populate_table()
        self.on_selection_changed()

    def populate_table(self):
        mode = self._mode()
        if mode == "short":
            columns = ["id", "score", "access", "content", "keywords", "last_accessed"]
        elif mode == "long":
            columns = ["id", "category", "score", "content", "reason", "promoted_at"]
        else:
            columns = ["id", "field", "confidence", "evidence", "claim", "updated_at"]

        self.table.clear()
        self.table.setColumnCount(len(columns))
        self.table.setHorizontalHeaderLabels(columns)
        self.table.setRowCount(len(self.current_rows))

        for r, row in enumerate(self.current_rows):
            if mode == "short":
                values = [
                    row.get("id"), row.get("importance_score"), row.get("access_count"),
                    row.get("content"), row.get("keywords"), row.get("last_accessed_at"),
                ]
            elif mode == "long":
                values = [
                    row.get("id"), row.get("category"), row.get("importance_score"),
                    row.get("content"), row.get("promote_reason"), row.get("promoted_at"),
                ]
            else:
                values = [
                    row.get("id"), row.get("field_name"), row.get("confidence"),
                    row.get("evidence_count"), row.get("claim"), row.get("updated_at"),
                ]
            for c, value in enumerate(values):
                item = QTableWidgetItem("" if value is None else str(value))
                item.setData(Qt.UserRole, row)
                item.setForeground(QBrush(QColor("#EAE5F2")))
                if r % 2:
                    item.setBackground(QBrush(QColor("#231B32")))
                else:
                    item.setBackground(QBrush(QColor("#1A1525")))
                self.table.setItem(r, c, item)

        self.table.resizeColumnsToContents()
        if self.table.columnCount() > 3:
            stretch_col = 4 if mode == "profile" else 3
            self.table.horizontalHeader().setSectionResizeMode(stretch_col, QHeaderView.Stretch)
        self.promote_btn.setEnabled(mode == "short")
        self.keywords_input.setEnabled(mode != "profile")
        self.category_input.setEnabled(mode != "short")

    def selected_row(self):
        indexes = self.table.selectionModel().selectedRows()
        if not indexes:
            return None
        r = indexes[0].row()
        if 0 <= r < len(self.current_rows):
            return self.current_rows[r]
        return None

    def on_selection_changed(self):
        row = self.selected_row()
        mode = self._mode()
        if not row:
            self.selected_label.setText("未选择")
            self.detail_edit.setPlainText("")
            self.category_input.setText("")
            self.keywords_input.setText("")
            return
        if mode == "short":
            self.selected_label.setText(f"短期记忆 #{row.get('id')}")
            self.detail_edit.setPlainText(str(row.get("content") or ""))
            self.category_input.setText("")
            self.keywords_input.setText(str(row.get("keywords") or ""))
        elif mode == "long":
            self.selected_label.setText(f"长期记忆 #{row.get('id')}")
            self.detail_edit.setPlainText(str(row.get("content") or ""))
            self.category_input.setText(str(row.get("category") or ""))
            self.keywords_input.setText(str(row.get("keywords") or ""))
        else:
            self.selected_label.setText(f"用户画像 claim #{row.get('id')}")
            self.detail_edit.setPlainText(str(row.get("claim") or ""))
            self.category_input.setText(str(row.get("field_name") or ""))
            self.keywords_input.setText("")

    def save_selected(self):
        row = self.selected_row()
        if not row:
            return
        mode = self._mode()
        content = self.detail_edit.toPlainText().strip()
        category = self.category_input.text().strip()
        keywords = self.keywords_input.text().strip()
        if mode == "profile" and _is_bad_profile_claim(category, content):
            QMessageBox.warning(self, "画像内容不合适", "这条内容更像应用/角色设定或过度推断，不适合作为用户画像。")
            return
        conn = None
        try:
            conn = self._connect_db()
            with conn.cursor() as cursor:
                if mode == "short":
                    cursor.execute(
                        "UPDATE user_memory SET content = %s, keywords = %s WHERE id = %s",
                        (content, keywords, row.get("id")),
                    )
                    chroma_ok_after_commit = (
                        row.get("id"),
                        content,
                        float(row.get("importance_score") or 0),
                    )
                elif mode == "long":
                    cursor.execute(
                        "UPDATE long_term_memory SET content = %s, keywords = %s, category = %s WHERE id = %s",
                        (content, keywords, category or _profile_bucket_for_memory(content), row.get("id")),
                    )
                else:
                    field = _normalize_profile_category(category or row.get("field_name"), content)
                    claim = _clean_profile_claim_text(content)
                    cursor.execute(
                        "UPDATE profile_claim SET field_name = %s, claim = %s, confidence = %s, updated_at = NOW() "
                        "WHERE id = %s",
                        (field, claim, 1.0, row.get("id")),
                    )
                    _refresh_user_profile_from_claims(cursor, clear_when_empty=True)
            conn.commit()
            if mode == "short" and 'chroma_ok_after_commit' in locals():
                mid, synced_content, synced_score = chroma_ok_after_commit
                if not sync_short_memory_to_chroma(mid, synced_content, synced_score):
                    QMessageBox.warning(self, "Chroma 同步失败", "MySQL 已保存，但 Chroma 同步失败。")
            if mode == "long":
                delete_profile_evidence_for_source("long_term_memory_edit", row.get("id"))
                schedule_profile_refine(
                    content,
                    source_type="long_term_memory",
                    source_ref=row.get("id"),
                    category_hint=category or _profile_bucket_for_memory(content),
                )
            QMessageBox.information(self, "已保存", "修改已经保存。")
            self.refresh()
        except Exception as e:
            QMessageBox.warning(self, "保存失败", f"保存失败：{e}")
        finally:
            try:
                if conn:
                    conn.close()
            except Exception:
                pass

    def clear_form(self):
        self.table.clearSelection()
        self.selected_label.setText("新建")
        self.detail_edit.setPlainText("")
        self.category_input.setText("")
        self.keywords_input.setText("")

    def add_new(self):
        mode = self._mode()
        content = self.detail_edit.toPlainText().strip()
        category = self.category_input.text().strip()
        keywords = self.keywords_input.text().strip()
        if not content:
            QMessageBox.warning(self, "内容为空", "请先在下方文本框填写要添加的记忆内容。")
            return
        if mode == "profile" and _is_bad_profile_claim(category, content):
            QMessageBox.warning(self, "画像内容不合适", "这条内容更像应用/角色设定或过度推断，不适合作为用户画像。")
            return

        conn = None
        saved_id = None
        try:
            conn = self._connect_db()
            with conn.cursor() as cursor:
                if mode == "short":
                    initial_score = _initial_importance_for_memory(content)
                    cursor.execute(
                        "INSERT INTO user_memory (content, keywords, importance_score) VALUES (%s, %s, %s)",
                        (content, keywords, initial_score),
                    )
                    saved_id = cursor.lastrowid
                elif mode == "long":
                    long_category = category or _profile_bucket_for_memory(content)
                    cursor.execute(
                        "INSERT INTO long_term_memory "
                        "(source_memory_id, content, content_hash, keywords, category, importance_score, promote_reason) "
                        "VALUES (NULL, %s, SHA2(%s, 256), %s, %s, %s, 'manual')",
                        (content, content, keywords, long_category, MEM_PROMOTE_TO_LONG_TERM_SCORE),
                    )
                    saved_id = cursor.lastrowid
                else:
                    field = _normalize_profile_category(category, content) if category else _profile_bucket_for_memory(content)
                    claim = _clean_profile_claim_text(content)
                    cursor.execute(
                        "INSERT INTO profile_claim "
                        "(field_name, claim, confidence, evidence_count, evidence_ids) "
                        "VALUES (%s, %s, 1.0, 1, NULL)",
                        (field, claim),
                    )
                    saved_id = cursor.lastrowid
                    _refresh_user_profile_from_claims(cursor, clear_when_empty=True)
            conn.commit()

            if mode == "short" and saved_id:
                if not sync_short_memory_to_chroma(
                    saved_id,
                    content,
                    _initial_importance_for_memory(content),
                ):
                    QMessageBox.warning(
                        self,
                        "Chroma 同步失败",
                        "短期记忆已写入 MySQL，但写入 Chroma 失败。",
                    )
            elif mode == "long":
                schedule_profile_refine(
                    content,
                    source_type="long_term_memory",
                    source_ref=saved_id,
                    category_hint=category or _profile_bucket_for_memory(content),
                )

            QMessageBox.information(self, "已添加", "新项已经添加。")
            self.refresh()
        except Exception as e:
            QMessageBox.warning(self, "添加失败", f"添加失败：{e}")
        finally:
            try:
                if conn:
                    conn.close()
            except Exception:
                pass

    def delete_selected(self):
        row = self.selected_row()
        if not row:
            return
        mode = self._mode()
        reply = QMessageBox.question(
            self, "确认删除", "确定删除选中的记忆/画像字段吗？",
            QMessageBox.Yes | QMessageBox.No
        )
        if reply != QMessageBox.Yes:
            return
        conn = None
        try:
            conn = self._connect_db()
            with conn.cursor() as cursor:
                if mode == "short":
                    mid = int(row.get("id"))
                    cursor.execute("DELETE FROM user_memory WHERE id = %s", (mid,))
                    chroma_delete_after_commit = mid
                elif mode == "long":
                    long_id = row.get("id")
                    cursor.execute("DELETE FROM long_term_memory WHERE id = %s", (long_id,))
                else:
                    cursor.execute("DELETE FROM profile_claim WHERE id = %s", (row.get("id"),))
                    _refresh_user_profile_from_claims(cursor, clear_when_empty=True)
            conn.commit()
            if mode == "short" and 'chroma_delete_after_commit' in locals():
                if not delete_short_memory_from_chroma(chroma_delete_after_commit):
                    QMessageBox.warning(self, "Chroma 同步失败", "MySQL 已删除，但 Chroma 删除失败。")
            if mode == "long":
                delete_profile_evidence_for_source("long_term_memory", long_id)
                delete_profile_evidence_for_source("long_term_memory_edit", long_id)
                refresh_user_profile_from_long_term(force=True)
            self.refresh()
        except Exception as e:
            QMessageBox.warning(self, "删除失败", f"删除失败：{e}")
        finally:
            try:
                if conn:
                    conn.close()
            except Exception:
                pass

    def promote_selected(self):
        row = self.selected_row()
        if not row or self._mode() != "short":
            return
        ok = promote_memory_to_long_term(row.get("id"), reason="manual")
        if ok:
            QMessageBox.information(self, "已迁移", "已迁移为长期记忆，并从短期记忆中移除。")
        else:
            QMessageBox.information(self, "未迁移", "没有迁移：可能已经迁移过，或短期记忆不存在。")
        self.refresh()

    def refresh_profile(self):
        refresh_user_profile_from_long_term(force=True)
        if self._mode() == "profile":
            self.refresh()
        QMessageBox.information(self, "已刷新", "用户画像已根据长期记忆重新聚合。")

    def refine_pending_memories(self):
        if not _profile_refiner_enabled():
            QMessageBox.information(
                self,
                "画像精炼未启用",
                "请先在设置里填写用户画像精炼模型的 API Key；未启用时不会消耗 tokens。",
            )
            return
        schedule_refine_unprocessed_long_term_memories(limit=10)
        QMessageBox.information(self, "已开始", "已在后台开始精炼最多 10 条未处理长期记忆。")


class FocusTimerWindow(QWidget):
    def __init__(self, pet):
        super().__init__()
        self.pet = pet
        self.is_collapsed = False
        self._expanded_geometry = None
        self._dragging = False
        self._drag_offset = None
        self.setWindowTitle("专注定时器")
        self.setWindowFlags(Qt.Window | Qt.WindowStaysOnTopHint)
        apply_dark_window_chrome(self)
        self.resize(320, 210)
        self.setStyleSheet("""
            QWidget {
                background: transparent;
                color: #EAE5F2;
                font-family: 'Segoe UI', 'Microsoft YaHei', sans-serif;
            }
            QFrame#focusTimerCard {
                background: #1A1525;
                border: 1px solid #3D2E55;
                border-radius: 20px;
            }
            QLabel#timerDisplay {
                background: transparent;
                color: #EAE5F2;
                border: none;
                font-size: 38px;
                font-weight: 600;
                padding: 6px 0 2px 0;
            }
            QLabel#fieldLabel {
                color: #B8ADC9;
                font-size: 12px;
            }
            QLabel#timerTitle {
                color: #B8ADC9;
                font-size: 12px;
                font-weight: 500;
            }
            QSpinBox {
                background: #231B32;
                color: #EAE5F2;
                border: 1px solid #4E3C6B;
                border-radius: 12px;
                padding: 5px 8px;
                font-size: 14px;
            }
            QSpinBox:focus {
                border-color: #B886F8;
                background: #20192D;
            }
            QPushButton {
                border: 1px solid #3D2E55;
                border-radius: 16px;
                padding: 8px 12px;
                background: #231B32;
                color: #EAE5F2;
                font-size: 14px;
            }
            QPushButton:hover { background: #2A203B; border-color: #B886F8; }
            QPushButton#primaryButton {
                background: #3D2E55;
                color: #EAE5F2;
                border-color: #B886F8;
                font-weight: 600;
            }
            QPushButton#primaryButton:hover { background: #4E3C6B; border-color: #A08CD2; }
            QPushButton#timerCloseButton {
                background: transparent;
                color: #B8ADC9;
                border: none;
                border-radius: 12px;
                padding: 0;
                font-size: 14px;
                font-weight: bold;
            }
            QPushButton#timerCloseButton:hover {
                background: #2A203B;
                color: #EAE5F2;
            }
        """)
        self._build_ui()
        self.refresh_state()

    def showEvent(self, event):
        super().showEvent(event)

    def _build_ui(self):
        self.root_layout = QVBoxLayout(self)
        self.root_layout.setContentsMargins(10, 10, 10, 10)
        self.root_layout.setSpacing(0)

        self.timer_card = QFrame()
        self.timer_card.setObjectName("focusTimerCard")
        self.card_layout = QVBoxLayout(self.timer_card)
        self.card_layout.setContentsMargins(16, 12, 16, 14)
        self.card_layout.setSpacing(8)

        self.title_panel = QWidget()
        title_layout = QHBoxLayout(self.title_panel)
        title_layout.setContentsMargins(0, 0, 0, 0)
        title_layout.setSpacing(8)
        timer_title = QLabel("专注计时")
        timer_title.setObjectName("timerTitle")
        title_layout.addWidget(timer_title)
        title_layout.addStretch()
        self.close_btn = QPushButton("×")
        self.close_btn.setObjectName("timerCloseButton")
        self.close_btn.setFixedSize(24, 24)
        self.close_btn.setToolTip("关闭计时器窗口")
        self.close_btn.clicked.connect(self.close)
        title_layout.addWidget(self.close_btn)
        self.card_layout.addWidget(self.title_panel)

        self.display_label = QLabel("25:00")
        self.display_label.setObjectName("timerDisplay")
        self.display_label.setAlignment(Qt.AlignCenter)
        self.display_label.setMinimumHeight(64)
        self.display_label.mousePressEvent = self._on_display_label_mouse_press
        self.display_label.mouseMoveEvent = self._on_display_label_mouse_move
        self.display_label.mouseReleaseEvent = self._on_display_label_mouse_release
        self.card_layout.addWidget(self.display_label)

        self.spin_panel = QWidget()
        spin_layout = QHBoxLayout(self.spin_panel)
        spin_layout.setContentsMargins(0, 0, 0, 0)
        spin_layout.setSpacing(10)
        self.hour_spin = self._make_spin(0, 23, 0)
        self.minute_spin = self._make_spin(0, 59, 25)
        self.second_spin = self._make_spin(0, 59, 0)
        for label, spin in (("小时", self.hour_spin), ("分钟", self.minute_spin), ("秒", self.second_spin)):
            field = QWidget()
            box = QVBoxLayout(field)
            box.setContentsMargins(0, 0, 0, 0)
            box.setSpacing(4)
            title = QLabel(label)
            title.setObjectName("fieldLabel")
            title.setAlignment(Qt.AlignCenter)
            box.addWidget(title)
            box.addWidget(spin)
            spin_layout.addWidget(field, 1)
        self.card_layout.addWidget(self.spin_panel)

        self.button_panel = QWidget()
        btn_layout = QHBoxLayout(self.button_panel)
        btn_layout.setContentsMargins(0, 0, 0, 0)
        btn_layout.setSpacing(8)
        self.start_pause_btn = QPushButton("开始")
        self.start_pause_btn.setObjectName("primaryButton")
        self.start_pause_btn.clicked.connect(self.on_start_pause)
        self.reset_btn = QPushButton("重置")
        self.reset_btn.clicked.connect(self.on_reset)
        self.collapse_btn = QPushButton("收起")
        self.collapse_btn.clicked.connect(self.collapse_to_badge)
        btn_layout.addWidget(self.start_pause_btn)
        btn_layout.addWidget(self.reset_btn)
        btn_layout.addWidget(self.collapse_btn)
        self.card_layout.addWidget(self.button_panel)
        self.root_layout.addWidget(self.timer_card)

    def _make_spin(self, min_value, max_value, value):
        spin = QSpinBox()
        spin.setRange(min_value, max_value)
        spin.setValue(value)
        spin.setAlignment(Qt.AlignCenter)
        spin.setFixedHeight(34)
        spin.setButtonSymbols(QSpinBox.UpDownArrows)
        spin.valueChanged.connect(self.refresh_idle_display)
        return spin

    def selected_seconds(self):
        return (
            self.hour_spin.value() * 3600
            + self.minute_spin.value() * 60
            + self.second_spin.value()
        )

    def set_spin_seconds(self, seconds):
        seconds = int(max(0, seconds))
        self.hour_spin.blockSignals(True)
        self.minute_spin.blockSignals(True)
        self.second_spin.blockSignals(True)
        self.hour_spin.setValue(min(23, seconds // 3600))
        seconds %= 3600
        self.minute_spin.setValue(seconds // 60)
        self.second_spin.setValue(seconds % 60)
        self.hour_spin.blockSignals(False)
        self.minute_spin.blockSignals(False)
        self.second_spin.blockSignals(False)

    def refresh_idle_display(self):
        if self.pet and self.pet.is_focus_timer_active():
            return
        self.display_label.setText(self._format_clock(max(1, self.selected_seconds())))

    def refresh_state(self):
        remaining = self.pet.get_focus_timer_remaining_seconds() if self.pet else 0
        if self.pet and self.pet.focus_timer_paused:
            self.display_label.setText(self._format_clock(remaining))
            self.start_pause_btn.setText("继续")
            self.reset_btn.setText("重置")
            return
        if self.pet and self.pet.focus_timer_end_at:
            self.display_label.setText(self._format_clock(remaining))
            self.start_pause_btn.setText("暂停")
            self.reset_btn.setText("重置")
            return
        self.start_pause_btn.setText("开始")
        self.reset_btn.setText("重置")
        self.refresh_idle_display()

    def collapse_to_badge(self):
        if self.is_collapsed:
            return
        self._expanded_geometry = self.geometry()
        self.is_collapsed = True
        self.spin_panel.hide()
        self.button_panel.hide()
        self.title_panel.hide()
        self.root_layout.setContentsMargins(0, 0, 0, 0)
        self.card_layout.setContentsMargins(8, 6, 8, 6)
        self.card_layout.setSpacing(0)
        self.display_label.setCursor(Qt.PointingHandCursor)
        self.display_label.setMinimumHeight(44)
        self.display_label.setMaximumHeight(44)
        self.display_label.setStyleSheet("""
            QLabel#timerDisplay {
                background: transparent;
                color: #EAE5F2;
                font-size: 18px;
                font-weight: 600;
                padding: 0;
            }
        """)
        self.setMinimumSize(112, 56)
        self.setMaximumSize(112, 56)
        self.resize(112, 56)
        screen = QApplication.primaryScreen()
        if screen:
            rect = screen.availableGeometry()
            self.move(rect.right() - self.width() - 18, rect.top() + 18)
        self.show()
        self.raise_()
        self.refresh_state()

    def expand_from_badge(self):
        if not self.is_collapsed:
            return
        self.is_collapsed = False
        self.setMaximumSize(16777215, 16777215)
        self.setMinimumSize(0, 0)
        self.title_panel.show()
        self.spin_panel.show()
        self.button_panel.show()
        self.root_layout.setContentsMargins(10, 10, 10, 10)
        self.card_layout.setContentsMargins(16, 12, 16, 14)
        self.card_layout.setSpacing(8)
        self.display_label.setCursor(Qt.ArrowCursor)
        self.display_label.setMinimumHeight(64)
        self.display_label.setMaximumHeight(16777215)
        self.display_label.setStyleSheet("")
        if self._expanded_geometry:
            self.setGeometry(self._expanded_geometry)
        else:
            self.resize(320, 210)
        self.show()
        self.raise_()
        self.activateWindow()
        self.refresh_state()

    def mousePressEvent(self, event):
        if self.is_collapsed and event.button() == Qt.LeftButton:
            self.expand_from_badge()
            return
        if not self.is_collapsed and _begin_window_drag(self, event, max_y=84):
            return
        super().mousePressEvent(event)

    def mouseMoveEvent(self, event):
        if _continue_window_drag(self, event):
            return
        super().mouseMoveEvent(event)

    def mouseReleaseEvent(self, event):
        _end_window_drag(self)
        super().mouseReleaseEvent(event)

    def _on_display_label_mouse_press(self, event):
        if self.is_collapsed and event.button() == Qt.LeftButton:
            self.expand_from_badge()
            event.accept()
            return
        if not self.is_collapsed and event.button() == Qt.LeftButton:
            self._dragging = True
            self._drag_offset = event.globalPos() - self.frameGeometry().topLeft()
            event.accept()
            return
        QLabel.mousePressEvent(self.display_label, event)

    def _on_display_label_mouse_move(self, event):
        if _continue_window_drag(self, event):
            return
        QLabel.mouseMoveEvent(self.display_label, event)

    def _on_display_label_mouse_release(self, event):
        _end_window_drag(self)
        QLabel.mouseReleaseEvent(self.display_label, event)

    def on_start_pause(self):
        if not self.pet:
            return
        if self.pet.focus_timer_end_at:
            self.pet.pause_focus_timer()
            return
        if self.pet.focus_timer_paused:
            self.pet.resume_focus_timer()
            return
        seconds = self.selected_seconds()
        if seconds <= 0:
            seconds = 25 * 60
            self.set_spin_seconds(seconds)
        self.pet.start_focus_timer(seconds, label="专注")

    def on_reset(self):
        if self.pet:
            self.pet.cancel_focus_timer(show_message=False)
        self.set_spin_seconds(25 * 60)
        self.display_label.setText("25:00")
        self.start_pause_btn.setText("开始")

    @staticmethod
    def _format_clock(seconds):
        seconds = int(max(0, seconds))
        h = seconds // 3600
        m = (seconds % 3600) // 60
        s = seconds % 60
        if h:
            return f"{h:02d}:{m:02d}:{s:02d}"
        return f"{m:02d}:{s:02d}"


class DesktopPet(QWidget):
    def __init__(self):
        super().__init__()
        # 初始化数据库及记忆衰减
        init_db()
        daily_decay_memory()
        refresh_user_profile_from_long_term(force=True)
        schedule_chroma_sync_repair()
        
        # 每天定时再执行一次遗忘整理 (86400000 毫秒 = 24小时)
        self.decay_timer = QTimer()
        self.decay_timer.timeout.connect(daily_decay_memory)
        self.decay_timer.start(1000 * 60 * 60 * 24)
        
        self.todo_window = None
        self.chat_window = None
        self.memory_window = None
        self.learning_label_window = None
        self.rss_manager_window = None
        self.focus_timer_window = None
        self.offset = QPoint()
        self.current_frame = 0
        self.is_happy = False
        self.happy_timer = 0
        self.proactive_care_thread = None
        self.rss_recommend_thread = None
        self.proactive_generation_seq = 0
        self.proactive_auto_enabled = True
        self.last_user_interaction_at = None
        self.last_proactive_at = None
        self.proactive_today = datetime.date.today()
        self.proactive_count_today = 0
        self.pending_proactive_messages = []
        self.rss_recommend_today = datetime.date.today()
        self.rss_recommend_count_today = 0
        self.last_rss_recommend_at = None
        self.awaiting_proactive_reply = False
        self.awaiting_proactive_reply_event_id = ""
        self.awaiting_proactive_reply_since = None
        self.last_proactive_silence_log_at = None
        self.focus_timer_end_at = None
        self.focus_timer_label = "专注"
        self.focus_timer_total_seconds = 0
        self.focus_timer_remaining_seconds = 0
        self.focus_timer_paused = False
        self.local_care_enabled = True
        self.rsshub_starter_thread = None
        self.todo_deadline_reminded = {}

        try:
            configure_local_rsshub_base()
            self.rsshub_starter_thread = start_local_rsshub_background()
        except Exception as e:
            print(f"[RSSHub] 本地 RSSHub 启动调度失败：{e}")
        
        # 主动关怀：根据上下文低频发聊天消息，和本地生活提醒分开。
        self.chat_timer = QTimer(self)
        self.chat_timer.timeout.connect(self.proactive_companion_tick)
        self.chat_timer.start(1000 * 60 * 10)  # 10分钟检查一次，实际发言会被冷却限制

        # 本地生活提醒：启动即运行，只弹气泡，不走 LLM，不写聊天历史。
        self.local_bubble_timer = QTimer(self)
        self.local_bubble_timer.timeout.connect(self.local_bubble_tick)
        self.local_bubble_timer.start(1000 * 60 * 30)
        
        self.sit_timer = QTimer(self)
        self.sit_timer.timeout.connect(self.remind_stand_up)
        self.sit_timer.start(1000 * 60 * 45)
        
        self.water_timer = QTimer(self)
        self.water_timer.timeout.connect(self.remind_drink_water)
        self.water_timer.start(1000 * 60 * 30)

        self.todo_deadline_timer = QTimer(self)
        self.todo_deadline_timer.timeout.connect(self.todo_deadline_tick)
        self.todo_deadline_timer.start(1000 * 60 * 5)

        rss_check_minutes = self._config_int("rss_recommender.background_check_minutes", 60, min_value=15, max_value=1440)
        self.rss_content_timer = QTimer(self)
        self.rss_content_timer.timeout.connect(self.rss_content_recommendation_tick)
        self.rss_content_timer.start(1000 * 60 * rss_check_minutes)

        rss_cleanup_hours = self._config_int("rss_recommender.cache_cleanup_interval_hours", 12, min_value=1, max_value=168)
        self.rss_cleanup_timer = QTimer(self)
        self.rss_cleanup_timer.timeout.connect(self.cleanup_rss_cache_tick)
        self.rss_cleanup_timer.start(1000 * 60 * 60 * rss_cleanup_hours)

        self.init_ui()
        self.init_animation()
        self.check_special_day()
        self.schedule_startup_local_care()
        QTimer.singleShot(1000 * 25, self.warmup_rss_content)
        QTimer.singleShot(1000 * 45, self.cleanup_rss_cache_tick)
        QTimer.singleShot(1000 * 90, self.rss_content_startup_recommendation_tick)
        schedule_auto_label_batch(reason="startup_backlog", delay_seconds=60)

    def init_ui(self):
        self.setWindowFlags(
            Qt.FramelessWindowHint |
            Qt.WindowStaysOnTopHint |
            Qt.Tool
        )
        self.setAttribute(Qt.WA_TranslucentBackground)
        self.setFixedSize(180, 200)

        self.pet_label = QLabel(self)
        self.pet_label.setAlignment(Qt.AlignCenter)
        self.pet_label.setGeometry(0, 30, 180, 170)

        # 气泡容器
        self.bubble_container = QWidget(self)
        self.bubble_container.setGeometry(0, 0, 180, 50)
        self.bubble_container.hide()
        
        # 气泡背景和文字
        self.bubble = QLabel(self.bubble_container)
        self.bubble.setAlignment(Qt.AlignLeft | Qt.AlignVCenter)
        self.bubble.setWordWrap(True)
        self.bubble.setStyleSheet("""
            QLabel {
                background-color: rgba(17, 17, 27, 238);
                border: 1px solid #B886F8;
                border-radius: 16px;
                padding: 6px 25px 6px 12px;
                font-size: 13px;
                color: #EAE5F2;
                font-family: 'Segoe UI', 'Microsoft YaHei', sans-serif;
            }
        """)
        self.bubble.setGeometry(0, 0, 180, 50)
        
        # 气泡关闭按钮
        self.close_bubble_btn = QPushButton("✕", self.bubble_container)
        self.close_bubble_btn.setGeometry(160, 5, 15, 15)
        self.close_bubble_btn.setStyleSheet("""
            QPushButton {
                background-color: transparent;
                color: #B886F8;
                border: none;
                font-size: 12px;
                font-weight: bold;
            }
            QPushButton:hover {
                color: #EAE5F2;
            }
        """)
        self.close_bubble_btn.clicked.connect(self.hide_bubble)
        
        # 打字机效果相关属性
        self.typewriter_timer = QTimer()
        self.typewriter_timer.timeout.connect(self.type_next_char)
        self.full_text = ""
        self.current_text = ""
        self.char_index = 0
        self.bubble_hide_timer = QTimer()
        self.bubble_hide_timer.timeout.connect(self.hide_bubble)
        self.bubble_hide_timer.setSingleShot(True)

        self.focus_timer_badge = QLabel(self)
        self.focus_timer_badge.setAlignment(Qt.AlignCenter)
        self.focus_timer_badge.setGeometry(96, 4, 78, 28)
        self.focus_timer_badge.setStyleSheet("""
            QLabel {
                background-color: rgba(7, 9, 13, 238);
                border: 1px solid #4E3C6B;
                border-radius: 16px;
                color: #EAE5F2;
                font-size: 12px;
                font-weight: bold;
                font-family: 'Segoe UI', 'Microsoft YaHei', sans-serif;
            }
        """)
        self.focus_timer_badge.hide()

        self.focus_countdown_timer = QTimer(self)
        self.focus_countdown_timer.timeout.connect(self.update_focus_timer_badge)

        screen_obj = QApplication.primaryScreen()
        if screen_obj:
            screen = screen_obj.geometry()
            self.move(screen.width() - 250, screen.height() - 300)
        else:
            self.move(100, 100)

    def init_animation(self):
        self.pet_movie = None
        gif_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "有珠.gif")
        if os.path.exists(gif_path):
            self.pet_movie = QMovie(gif_path)
            self.pet_movie.setScaledSize(QSize(180, 170))
            self.pet_label.setMovie(self.pet_movie)
            self.pet_movie.start()
        else:
            self.normal_frames = [
                self.create_pet_pixmap("😺", "ᓚᘏᗢ"),
                self.create_pet_pixmap("😺", "ᓚᘏᗢ ~"),
                self.create_pet_pixmap("😺", "ᓚᘏᗢ  ~"),
                self.create_pet_pixmap("😸", "ᓚᘏᗢ"),
            ]
            self.happy_frames = [
                self.create_pet_pixmap("😻", "♡"),
                self.create_pet_pixmap("🥰", "♡♡"),
                self.create_pet_pixmap("😻", "♡♡♡"),
            ]
            self.pet_label.setPixmap(self.normal_frames[0])

        self.anim_timer = QTimer()
        self.anim_timer.timeout.connect(self.update_animation)
        self.anim_timer.start(500)

    def create_pet_pixmap(self, emoji, decoration=""):
        pixmap = QPixmap(180, 170)
        pixmap.fill(Qt.transparent)

        painter = QPainter(pixmap)
        painter.setRenderHint(QPainter.Antialiasing)

        # 身体
        painter.setBrush(QColor(255, 200, 100, 200))
        painter.setPen(QColor(200, 150, 50))
        painter.drawEllipse(30, 30, 120, 120)

        # 耳朵
        painter.setBrush(QColor(255, 180, 80, 200))

        left_ear = QPolygon([
            QPoint(45, 45),
            QPoint(30, 5),
            QPoint(70, 35)
        ])
        painter.drawPolygon(left_ear)

        right_ear = QPolygon([
            QPoint(135, 45),
            QPoint(150, 5),
            QPoint(110, 35)
        ])
        painter.drawPolygon(right_ear)

        # 脸
        font = QFont("Segoe UI Emoji", 40)
        painter.setFont(font)
        painter.drawText(55, 115, emoji)

        # 装饰文字
        font2 = QFont("Arial", 14)
        painter.setFont(font2)
        painter.setPen(QColor(255, 100, 100))
        painter.drawText(50, 165, decoration)

        painter.end()
        return pixmap

    def update_animation(self):
        if self.pet_movie is not None:
            if self.is_happy:
                self.pet_movie.setSpeed(150)
                self.happy_timer -= 1
                if self.happy_timer <= 0:
                    self.is_happy = False
            else:
                self.pet_movie.setSpeed(100)
            return

        if self.is_happy:
            frames = self.happy_frames
            self.happy_timer -= 1
            if self.happy_timer <= 0:
                self.is_happy = False
        else:
            frames = self.normal_frames

        self.current_frame = (self.current_frame + 1) % len(frames)
        self.pet_label.setPixmap(frames[self.current_frame])

    def check_special_day(self):
        now = datetime.datetime.now()
        year, month, day = now.year, now.month, now.day
        lunar_date = ZhDate.from_datetime(now)
        # 定义节日和对应的问候语、皮肤
        # 默认生日设为 5月20日，你可以修改为自己的生日
        self.special_days = {
            (1, 1): {"greeting": "元旦快乐！新的一年也要加油哦！🎉", "skin": "有珠.gif"},
            (2, 14): {"greeting": "情人节快乐！今天也要开心呀~ 💖", "skin": "有珠.gif"},
            (9, 19): {"greeting": "生日快乐！愿你今天是最幸福的人！🎂🎁", "skin": "有珠.gif"}, # 修改为你的生日
            (9, 30): {"greeting": "今天是我的生日！谢谢你的祝福！🎉", "skin": "有珠.gif"},
            (10, 1): {"greeting": "国庆节快乐！好好休息一下吧！🇨🇳", "skin": "有珠.gif"},
            (12, 25): {"greeting": "圣诞快乐！收到礼物了吗？🎄🎅", "skin": "有珠.gif"},
            (2, 22): {"greeting": "今天是猫之日！可爱吗？", "skin": "猫猫有珠.gif"}
        }
        self.luner_special_days = {
            (1, 1): {"greeting": "新春快乐！祝你新的一年事事顺利健康快乐！🎇", "skin": "有珠.gif"},
            (1, 15): {"greeting": "元宵节快乐！今天去逛灯谜展了嘛！🎆", "skin": "有珠.gif"},
            (5, 5): {"greeting": "端午节快乐！今天吃的粽子味道咋样？粽子", "skin": "有珠.gif"},
            (8, 15): {"greeting": "中秋节快乐！月饼都吃了啥口味？🌙", "skin": "有珠.gif"},
        }
        
        today = (month, day)
        lunar_today = (lunar_date.lunar_month, lunar_date.lunar_day)
        info =None 
        if today in self.special_days:
            info = self.special_days[today]
        elif lunar_today in self.luner_special_days:
            info = self.luner_special_days[lunar_today]
        if info:
            self.show_bubble(info["greeting"], duration=8000)
            self.set_happy()
            skin_name = info.get("skin", "有珠.gif")
            skin_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), skin_name)
            if not os.path.exists(skin_path):
                skin_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "有珠.gif")
            if os.path.exists(skin_path) and self.pet_movie:
                self.pet_movie.stop()
                self.pet_movie.setFileName(skin_path)
                self.pet_movie.start()

    def observe_user_message(self, text):
        if self.awaiting_proactive_reply and self.awaiting_proactive_reply_event_id:
            try:
                event_id = self.awaiting_proactive_reply_event_id
                replied_after = None
                if self.awaiting_proactive_reply_since:
                    replied_after = int(
                        (datetime.datetime.now() - self.awaiting_proactive_reply_since).total_seconds()
                    )
                log_feedback_event(
                    event_id=event_id,
                    feedback_value=1,
                    feedback_scope="proactive_user_replied",
                    user_text=text,
                    assistant_text="",
                    extra={"user_replied_after_sec": replied_after},
                )
                schedule_auto_relabel_event(event_id, reason="proactive_user_replied")
            except Exception as e:
                print(f"[LearningLog] 主动关怀隐式反馈写入失败: {e}")
        self.last_user_interaction_at = datetime.datetime.now()
        self.awaiting_proactive_reply = False
        self.awaiting_proactive_reply_event_id = ""
        self.awaiting_proactive_reply_since = None
        QTimer.singleShot(1000 * 60 * 12, self.proactive_companion_tick)

    def _learning_app_state_snapshot(self):
        now = datetime.datetime.now()
        def _minutes_since(value):
            if not value:
                return None
            try:
                return round(max(0, (now - value).total_seconds()) / 60, 2)
            except Exception:
                return None
        return {
            "chat_visible": bool(self.chat_window is not None and self.chat_window.isVisible()),
            "bubble_visible": bool(getattr(self, "bubble_container", None) and self.bubble_container.isVisible()),
            "pending_proactive_count": len(self.pending_proactive_messages),
            "awaiting_proactive_reply": bool(self.awaiting_proactive_reply),
            "proactive_count_today": int(self.proactive_count_today),
            "proactive_daily_limit": 10,
            "minutes_since_last_proactive": _minutes_since(self.last_proactive_at),
            "minutes_since_last_user_interaction": _minutes_since(self.last_user_interaction_at),
            "focus_timer_running": bool(self.focus_timer_end_at and not self.focus_timer_paused),
            "focus_timer_remaining_seconds": int(self.focus_timer_remaining_seconds or 0),
            "local_care_enabled": bool(self.local_care_enabled),
        }

    def _learning_observation_delay_seconds(self):
        try:
            value = int(app_config.get("learning_labeler.observation_window_seconds", 900))
        except Exception:
            value = 900
        return max(60, min(value, 60 * 60))

    def _record_learning_state_observation(self, event_id, assistant_text="", scope="implicit_state_observation", delay_seconds=None, extra=None):
        event_id = str(event_id or "").strip()
        if not event_id:
            return
        try:
            payload = dict(extra or {})
            if scope == "proactive_followup_state":
                payload["user_replied_before_observation"] = not (
                    self.awaiting_proactive_reply
                    and self.awaiting_proactive_reply_event_id == event_id
                )
            system_state = collect_system_state(extra=self._learning_app_state_snapshot())
            log_implicit_state_observation(
                event_id=event_id,
                scope=scope,
                feedback_value=0,
                assistant_text=assistant_text,
                system_state=system_state,
                delay_seconds=delay_seconds,
                extra=payload,
            )
            schedule_auto_relabel_event(event_id, reason=scope)
        except Exception as e:
            print(f"[LearningLog] implicit state observation failed: {e}")

    def _schedule_learning_state_observation(self, event_id, assistant_text="", scope="implicit_state_observation", extra=None):
        event_id = str(event_id or "").strip()
        if not event_id:
            return
        delay_seconds = self._learning_observation_delay_seconds()
        QTimer.singleShot(
            delay_seconds * 1000,
            lambda eid=event_id, text=assistant_text, sc=scope, ds=delay_seconds, ex=dict(extra or {}):
                self._record_learning_state_observation(
                    eid,
                    assistant_text=text,
                    scope=sc,
                    delay_seconds=ds,
                    extra=ex,
                )
        )

    def _log_proactive_learning_event(self, *, text="", state="idle", selected_strategy="proactive_checkin", reason=""):
        try:
            try:
                recent_turns = conversation_history.get_turns()
            except Exception:
                recent_turns = []
            try:
                profile_snapshot = get_user_profile_prompt_context()
            except Exception:
                profile_snapshot = "无"
            event = build_interaction_event(
                user_text="[PROACTIVE_TRIGGER]",
                assistant_text=text,
                trigger_type="proactive_timer",
                trigger_source=state or "proactive",
                recent_turns=recent_turns,
                user_profile_snapshot=profile_snapshot,
                system_state_snapshot=collect_system_state(extra=self._learning_app_state_snapshot()),
                current_response_card=f"主动关怀触发；策略={selected_strategy}；原因={reason}",
                models={"proactive": app_config.get("ark.model_main", "") or "doubao-1-5-pro-32k-250115"},
            )
            event["strategy"]["selected"] = selected_strategy
            event["strategy"]["reason"] = reason
            event["strategy"]["recommendation_used"] = False
            log_interaction_event(event, enqueue_label=True)
            schedule_auto_label_batch(reason="proactive_event")
            return event
        except Exception as e:
            print(f"[LearningLog] 主动关怀样本写入失败: {e}")
            return {}

    def _maybe_log_proactive_silence(self, reason):
        now = datetime.datetime.now()
        if (
            self.last_proactive_silence_log_at
            and (now - self.last_proactive_silence_log_at).total_seconds() < 60 * 60
        ):
            return
        self.last_proactive_silence_log_at = now
        self._log_proactive_learning_event(
            text="",
            state="proactive_silence",
            selected_strategy="do_nothing",
            reason=reason,
        )

    def proactive_companion_tick(self):
        if not self.proactive_auto_enabled:
            return
        now = datetime.datetime.now()
        if 1 <= now.hour < 7:
            self._maybe_log_proactive_silence("night_quiet_hours")
            return
        if self.proactive_today != now.date():
            self.proactive_today = now.date()
            self.proactive_count_today = 0
        if self.proactive_count_today >= 10:
            self._maybe_log_proactive_silence("daily_limit_reached")
            return
        if self.awaiting_proactive_reply or self.pending_proactive_messages:
            self._maybe_log_proactive_silence("waiting_for_user_after_previous_proactive")
            return
        if self.last_proactive_at and (now - self.last_proactive_at).total_seconds() < 45 * 60:
            self._maybe_log_proactive_silence("cooldown")
            return
        if self.proactive_care_thread is not None and self.proactive_care_thread.isRunning():
            self._maybe_log_proactive_silence("generation_already_running")
            return
        self.proactive_generation_seq += 1
        self.proactive_care_thread = ProactiveCareThread(
            generation_id=self.proactive_generation_seq,
            started_at=now,
            last_user_interaction_at=self.last_user_interaction_at,
            pending_count=len(self.pending_proactive_messages),
            awaiting_reply=bool(self.awaiting_proactive_reply),
        )
        self.proactive_care_thread.finished_signal.connect(self.on_proactive_care_generated)
        self.proactive_care_thread.error_signal.connect(self.on_proactive_care_error)
        self.proactive_care_thread.start()

    def on_proactive_care_generated(self, text):
        thread = self.sender()
        if thread is not self.proactive_care_thread:
            print("[Proactive] 丢弃旧主动关怀：生成线程已不是当前线程")
            return
        if not self._is_proactive_generation_fresh(thread):
            print("[Proactive] 丢弃旧主动关怀：生成期间用户状态已变化")
            return
        text = str(text or "").strip()
        if not text:
            return
        self.send_proactive_message(text, state="llm_context")

    def _is_proactive_generation_fresh(self, thread):
        if not thread:
            return False
        if getattr(thread, "generation_id", None) != self.proactive_generation_seq:
            return False
        started_at = getattr(thread, "started_at", None)
        if self.last_user_interaction_at and started_at and self.last_user_interaction_at >= started_at:
            return False
        if len(self.pending_proactive_messages) != getattr(thread, "pending_count", 0):
            return False
        if bool(self.awaiting_proactive_reply) != bool(getattr(thread, "awaiting_reply", False)):
            return False
        return True

    def on_proactive_care_error(self, error):
        print(f"[Proactive] 主动关怀生成失败: {error}")

    def local_bubble_tick(self):
        if not self.local_care_enabled:
            return
        if self.bubble_container.isVisible():
            return
        hour = datetime.datetime.now().hour
        if 1 <= hour < 7:
            return
        messages = [
            "记得喝水。别等到口渴才想起来。",
            "坐久了就站起来走两步。",
            "眼睛离屏幕远一点。看久了会累。",
            "如果在学习，就先做完下一小段。",
            "肩膀放松。别一直绷着。",
            "有事要记就直接告诉我，我会写进待办。",
        ]
        self.show_local_care_bubble(random.choice(messages), duration=6000)

    @staticmethod
    def _parse_todo_due_datetime(raw):
        text = str(raw or "").strip()
        if not text or "无" in text or text.lower() in {"none", "null"}:
            return None
        text = text.replace("/", "-").replace("T", " ")
        text = re.sub(r"\s+", " ", text)
        date_only = re.fullmatch(r"\d{4}-\d{1,2}-\d{1,2}", text)
        if date_only:
            try:
                return datetime.datetime.strptime(text, "%Y-%m-%d").replace(hour=23, minute=59, second=59)
            except Exception:
                return None
        match = re.search(r"(\d{4}-\d{1,2}-\d{1,2})\s+(\d{1,2}:\d{1,2}(?::\d{1,2})?)", text)
        if match:
            text = f"{match.group(1)} {match.group(2)}"
        for fmt in ("%Y-%m-%d %H:%M:%S", "%Y-%m-%d %H:%M"):
            try:
                return datetime.datetime.strptime(text, fmt)
            except Exception:
                pass
        return None

    @staticmethod
    def _todo_deadline_tier(seconds_left):
        if seconds_left < -10 * 60:
            return "overdue"
        if 0 <= seconds_left <= 60 * 60:
            return "1h"
        if 0 <= seconds_left <= 6 * 60 * 60:
            return "6h"
        if 0 <= seconds_left <= 24 * 60 * 60:
            return "24h"
        return ""

    @staticmethod
    def _format_deadline_left(seconds_left):
        if seconds_left < 0:
            minutes = int(abs(seconds_left) // 60)
            if minutes < 60:
                return f"已经过了 {minutes} 分钟"
            hours = minutes // 60
            return f"已经过了 {hours} 小时"
        minutes = int(max(1, seconds_left // 60))
        if minutes < 60:
            return f"还剩 {minutes} 分钟"
        hours = minutes // 60
        if hours < 24:
            return f"还剩 {hours} 小时"
        days = hours // 24
        return f"还剩 {days} 天"

    def todo_deadline_tick(self):
        if not getattr(self, "local_care_enabled", True):
            return
        if hasattr(self, "bubble_container") and self.bubble_container.isVisible():
            return
        try:
            items = todo_store().all()
        except Exception as e:
            print(f"[TodoReminder] 读取待办失败: {e}")
            return
        now = datetime.datetime.now()
        candidates = []
        active_ids = set()
        for item in items:
            if not isinstance(item, dict) or item.get("completed"):
                continue
            todo_id = str(item.get("id") or "")
            active_ids.add(todo_id)
            due_at = self._parse_todo_due_datetime(item.get("endtime"))
            if not due_at:
                continue
            seconds_left = (due_at - now).total_seconds()
            tier = self._todo_deadline_tier(seconds_left)
            if not tier:
                continue
            if self.todo_deadline_reminded.get(todo_id) == tier:
                continue
            candidates.append((seconds_left, tier, item, due_at))

        for todo_id in list(self.todo_deadline_reminded):
            if todo_id not in active_ids:
                self.todo_deadline_reminded.pop(todo_id, None)

        if not candidates:
            return
        candidates.sort(key=lambda row: row[0])
        seconds_left, tier, item, due_at = candidates[0]
        text = str(item.get("text") or "有一条待办").strip()
        if len(text) > 34:
            text = text[:34] + "..."
        left = self._format_deadline_left(seconds_left)
        if tier == "overdue":
            bubble = f"待办过截止时间了：{text}。\n{left}，别忘了处理。"
        else:
            bubble = f"待办快截止了：{text}。\n{left}，截止 {due_at.strftime('%m-%d %H:%M')}。"
        self.todo_deadline_reminded[str(item.get("id") or "")] = tier
        self.show_local_care_bubble(bubble, duration=8000)

    def schedule_startup_local_care(self):
        QTimer.singleShot(1000 * 20, self.startup_local_care_tick)
        QTimer.singleShot(1000 * 60 * 30, self.remind_drink_water)
        QTimer.singleShot(1000 * 60 * 45, self.remind_stand_up)
        QTimer.singleShot(1000 * 60 * 3, self.proactive_companion_tick)

    def warmup_rss_content(self):
        try:
            rss_content_runtime.warmup()
        except Exception as e:
            print("[RSS] warmup failed:", e)

    def cleanup_rss_cache_tick(self):
        try:
            rss_content_runtime.cleanup_cache(force=False)
        except Exception as e:
            print("[RSS] cache cleanup failed:", e)

    @staticmethod
    def _config_bool(key, default=False):
        value = app_config.get(key, default)
        if value is None or value == "":
            return bool(default)
        return str(value).strip().lower() not in {"0", "false", "no", "off"}

    @staticmethod
    def _config_int(key, default, min_value=0, max_value=None):
        try:
            value = int(app_config.get(key, default))
        except Exception:
            value = int(default)
        value = max(int(min_value), value)
        if max_value is not None:
            value = min(int(max_value), value)
        return value

    def rss_content_startup_recommendation_tick(self):
        self.rss_content_recommendation_tick(trigger_type="startup")

    def _rss_recommendation_allowed(self, trigger_type):
        if not self._config_bool("rss_recommender.enabled", True):
            return False, "rss_disabled"
        if trigger_type == "startup" and not self._config_bool("rss_recommender.startup_recommend_enabled", True):
            return False, "startup_recommend_disabled"
        if trigger_type != "startup" and not self._config_bool("rss_recommender.idle_recommend_enabled", True):
            return False, "idle_recommend_disabled"

        now = datetime.datetime.now()
        if self.rss_recommend_today != now.date():
            self.rss_recommend_today = now.date()
            self.rss_recommend_count_today = 0
        daily_limit = self._config_int("rss_recommender.active_recommend_daily_limit", 2, min_value=0, max_value=10)
        if daily_limit <= 0 or self.rss_recommend_count_today >= daily_limit:
            return False, "daily_limit"
        cooldown_hours = self._config_int("rss_recommender.active_recommend_cooldown_hours", 8, min_value=1, max_value=72)
        if self.last_rss_recommend_at and (now - self.last_rss_recommend_at).total_seconds() < cooldown_hours * 3600:
            return False, "cooldown"
        if self.awaiting_proactive_reply or self.pending_proactive_messages:
            return False, "waiting_for_user_after_previous_proactive"
        if self.rss_recommend_thread is not None and self.rss_recommend_thread.isRunning():
            return False, "rss_recommend_running"
        if self.proactive_care_thread is not None and self.proactive_care_thread.isRunning():
            return False, "proactive_generation_running"
        if trigger_type != "startup":
            try:
                state = collect_system_state()
                idle = state.get("idle") or {}
                if not bool(idle.get("is_idle")):
                    return False, "not_idle"
            except Exception:
                return False, "idle_state_unavailable"
        return True, "ok"

    def rss_content_recommendation_tick(self, trigger_type="proactive_timer"):
        allowed, reason = self._rss_recommendation_allowed(trigger_type)
        if not allowed:
            print(f"[RSS] skip proactive content recommendation: {reason}")
            return
        self.rss_recommend_thread = RSSContentRecommendThread(trigger_type=trigger_type)
        self.rss_recommend_thread.finished_signal.connect(self.on_rss_content_recommended)
        self.rss_recommend_thread.error_signal.connect(self.on_rss_content_recommend_error)
        self.rss_recommend_thread.start()

    def on_rss_content_recommended(self, payload):
        payload = payload or {}
        message = str(payload.get("message") or "").strip()
        decision = payload.get("decision") or {}
        if not message:
            print(f"[RSS] no proactive content recommendation: {decision.get('reason', '')}")
            return
        self.last_rss_recommend_at = datetime.datetime.now()
        self.rss_recommend_count_today += 1
        self.send_proactive_message(
            message,
            state="rss_content_recommendation",
            selected_strategy="external_content_recommendation",
            reason="rss_background_recommendation",
            external_content_decision=decision,
        )

    def on_rss_content_recommend_error(self, error):
        print(f"[RSS] proactive content recommendation failed: {error}")

    def startup_local_care_tick(self):
        self.show_local_care_bubble("我在。水放近一点，别一直盯着屏幕。", duration=6000)

    def show_local_care_bubble(self, text, duration=6000):
        if not self.local_care_enabled:
            return False
        hour = datetime.datetime.now().hour
        if 1 <= hour < 7:
            return False
        if self.bubble_container.isVisible():
            QTimer.singleShot(1000 * 25, lambda: self.show_local_care_bubble(text, duration))
            return False
        self.set_happy()
        self.show_bubble(text, duration=duration)
        return True

    def send_proactive_message(
            self, text, state="idle", selected_strategy="proactive_context_checkin",
            reason="llm_generated_active_care", external_content_decision=None):
        if not text:
            return
        self.last_proactive_at = datetime.datetime.now()
        self.proactive_count_today += 1
        self.awaiting_proactive_reply = True
        self.awaiting_proactive_reply_since = self.last_proactive_at
        learning_event = self._log_proactive_learning_event(
            text=text,
            state=state,
            selected_strategy=selected_strategy,
            reason=reason,
        )
        self.awaiting_proactive_reply_event_id = learning_event.get("event_id", "")
        if self.awaiting_proactive_reply_event_id:
            self._schedule_learning_state_observation(
                self.awaiting_proactive_reply_event_id,
                assistant_text=text,
                scope="proactive_followup_state",
                extra={
                    "proactive_state": state,
                    "user_replied_before_observation": False,
                    "pending_reply_event_id": self.awaiting_proactive_reply_event_id,
                },
            )
        try:
            conversation_history.add_turn("", text, source="proactive")
        except Exception as e:
            print(f"[Proactive] 写入对话历史失败: {e}")
        self.pending_proactive_messages.append({
            "text": text,
            "state": state,
            "time": self.last_proactive_at.isoformat(),
            "learning_event": learning_event,
            "external_content_decision": external_content_decision or {},
        })
        if self.chat_window is not None and self.chat_window.isVisible():
            old_feedback_context = getattr(self.chat_window, "last_user_message_for_feedback", "")
            self.chat_window.last_user_message_for_feedback = "[主动陪伴消息]"
            self.chat_window.add_message(
                text,
                is_user=False,
                learning_event=learning_event,
                external_content_decision=external_content_decision or {},
            )
            self.chat_window.last_user_message_for_feedback = old_feedback_context
            self.pending_proactive_messages.clear()
        if not self.bubble_container.isVisible():
            self.show_bubble("有事找你。看看聊天窗口。", duration=6000)

    def flush_pending_proactive_messages(self):
        if self.chat_window is None or not self.chat_window.isVisible():
            return
        for msg in self.pending_proactive_messages:
            old_feedback_context = getattr(self.chat_window, "last_user_message_for_feedback", "")
            self.chat_window.last_user_message_for_feedback = "[主动陪伴消息]"
            learning_event = msg.get("learning_event") or {}
            self.chat_window.add_message(
                msg.get("text", ""),
                is_user=False,
                learning_event=learning_event,
                external_content_decision=msg.get("external_content_decision") or {},
            )
            if learning_event.get("event_id"):
                try:
                    event_id = learning_event.get("event_id", "")
                    log_feedback_event(
                        event_id=event_id,
                        feedback_value=1,
                        feedback_scope="proactive_chat_opened",
                        user_text="[OPEN_CHAT_AFTER_PROACTIVE]",
                        assistant_text=msg.get("text", ""),
                        extra={"proactive_state": msg.get("state", "")},
                    )
                    schedule_auto_relabel_event(event_id, reason="proactive_chat_opened")
                except Exception as e:
                    print(f"[LearningLog] 主动关怀打开聊天反馈写入失败: {e}")
            self.chat_window.last_user_message_for_feedback = old_feedback_context
        self.pending_proactive_messages.clear()

    def random_chat(self):
        self.proactive_companion_tick()

    def remind_drink_water(self):
        self.show_local_care_bubble("喝水。现在就去倒一点，别等口渴。", duration=6500)
        
    def remind_stand_up(self):
        self.show_local_care_bubble("站起来走两步。肩膀和眼睛都该休息一下。", duration=6500)

    def show_bubble(self, text, duration=3000):
        self.typewriter_timer.stop()
        self.bubble_hide_timer.stop()
        
        self.full_text = text
        self.current_text = ""
        self.char_index = 0
        self.bubble.setText("")
        
        # 自适应气泡高度
        self.bubble.setText(text)
        self.bubble.adjustSize()
        height = max(50, self.bubble.height() + 10)
        self.bubble_container.setGeometry(0, 0, 180, height)
        self.bubble.setGeometry(0, 0, 180, height)
        self.bubble.setText("")
        
        self.bubble_container.setWindowOpacity(0.0)
        self.bubble_container.show()
        
        # 气泡弹出动画
        self.bubble_anim = QPropertyAnimation(self.bubble_container, b"windowOpacity")
        self.bubble_anim.setDuration(250)
        self.bubble_anim.setStartValue(0.0)
        self.bubble_anim.setEndValue(1.0)
        self.bubble_anim.start()
        
        if getattr(self, "focus_timer_badge", None) is not None and self.focus_timer_badge.isVisible():
            self.focus_timer_badge.raise_()
        
        # 存储气泡停留时间
        self.bubble_duration = duration
        self.typewriter_timer.start(50)  # 每个字 50ms 的速度

    def type_next_char(self):
        if self.char_index < len(self.full_text):
            self.current_text += self.full_text[self.char_index]
            self.bubble.setText(self.current_text)
            self.char_index += 1
        else:
            self.typewriter_timer.stop()
            if self.bubble_duration > 0:
                self.bubble_hide_timer.start(self.bubble_duration)
                
    def hide_bubble(self):
        self.typewriter_timer.stop()
        self.bubble_hide_timer.stop()
        
        # 气泡消失动画
        self.bubble_anim = QPropertyAnimation(self.bubble_container, b"windowOpacity")
        self.bubble_anim.setDuration(300)
        self.bubble_anim.setStartValue(1.0)
        self.bubble_anim.setEndValue(0.0)
        self.bubble_anim.finished.connect(self.bubble_container.hide)
        self.bubble_anim.start()

    def is_focus_timer_active(self):
        return bool(self.focus_timer_end_at or self.focus_timer_paused)

    def get_focus_timer_remaining_seconds(self):
        if self.focus_timer_paused:
            return int(max(0, self.focus_timer_remaining_seconds))
        if self.focus_timer_end_at:
            return int(max(0, (self.focus_timer_end_at - datetime.datetime.now()).total_seconds() + 0.9))
        return 0

    def _sync_focus_timer_window(self):
        if self.focus_timer_window is not None and self.focus_timer_window.isVisible():
            self.focus_timer_window.refresh_state()

    def _render_focus_timer_badge(self, remaining=None):
        if remaining is None:
            remaining = self.get_focus_timer_remaining_seconds()
        if self.focus_timer_paused:
            text = f"暂停 {format_focus_duration(remaining)}"
        else:
            text = format_focus_duration(remaining)
        self.focus_timer_badge.setText(text)
        self.focus_timer_badge.adjustSize()
        width = max(78, self.focus_timer_badge.width() + 14)
        self.focus_timer_badge.setGeometry(max(0, 176 - width), 4, width, 28)
        self.focus_timer_badge.show()
        self.focus_timer_badge.raise_()

    def start_focus_timer(self, seconds, label="专注"):
        seconds = int(max(5, min(int(seconds or 0), 24 * 60 * 60)))
        self.focus_timer_label = (label or "专注").strip()[:8] or "专注"
        self.focus_timer_total_seconds = seconds
        self.focus_timer_remaining_seconds = seconds
        self.focus_timer_paused = False
        self.focus_timer_end_at = datetime.datetime.now() + datetime.timedelta(seconds=seconds)
        self._render_focus_timer_badge(seconds)
        self.focus_countdown_timer.start(1000)
        self._sync_focus_timer_window()
        self.show_bubble(f"{self.focus_timer_label}定时开始：{format_focus_duration(seconds)}。", duration=3500)

    def pause_focus_timer(self):
        if not self.focus_timer_end_at:
            return False
        self.focus_timer_remaining_seconds = self.get_focus_timer_remaining_seconds()
        self.focus_timer_end_at = None
        self.focus_timer_paused = True
        self.focus_countdown_timer.stop()
        self._render_focus_timer_badge(self.focus_timer_remaining_seconds)
        self._sync_focus_timer_window()
        self.show_bubble("专注定时已暂停。", duration=2500)
        return True

    def resume_focus_timer(self):
        if not self.focus_timer_paused or self.focus_timer_remaining_seconds <= 0:
            return False
        self.focus_timer_paused = False
        self.focus_timer_end_at = datetime.datetime.now() + datetime.timedelta(seconds=self.focus_timer_remaining_seconds)
        self.focus_countdown_timer.start(1000)
        self._render_focus_timer_badge(self.focus_timer_remaining_seconds)
        self._sync_focus_timer_window()
        self.show_bubble("继续计时。", duration=2500)
        return True

    def cancel_focus_timer(self, show_message=True):
        if not self.is_focus_timer_active():
            if show_message:
                self.show_bubble("现在没有正在计时的专注定时。", duration=3000)
            return False
        self.focus_countdown_timer.stop()
        self.focus_timer_end_at = None
        self.focus_timer_paused = False
        self.focus_timer_remaining_seconds = 0
        self.focus_timer_total_seconds = 0
        self.focus_timer_badge.hide()
        self._sync_focus_timer_window()
        if show_message:
            self.show_bubble("专注定时取消了。", duration=3000)
        return True

    def update_focus_timer_badge(self):
        if not self.focus_timer_end_at:
            return
        remaining = self.get_focus_timer_remaining_seconds()
        self.focus_timer_remaining_seconds = remaining
        if remaining <= 0:
            self.finish_focus_timer()
            return
        self._render_focus_timer_badge(remaining)
        self._sync_focus_timer_window()

    def finish_focus_timer(self):
        label = self.focus_timer_label or "专注"
        self.focus_countdown_timer.stop()
        self.focus_timer_end_at = None
        self.focus_timer_paused = False
        self.focus_timer_remaining_seconds = 0
        self.focus_timer_total_seconds = 0
        self.focus_timer_badge.hide()
        self._sync_focus_timer_window()
        text = f"{label}定时结束了。先停一下，别把自己烧干。"
        self.show_bubble(text, duration=9000)
        if self.chat_window is not None and self.chat_window.isVisible():
            self.chat_window.add_message(f"⏱ {text}", is_user=False)

    def set_happy(self):
        self.is_happy = True
        self.happy_timer = 6
        if self.pet_movie is not None:
            self.pet_movie.setSpeed(150)

    def mousePressEvent(self, event):
        if event.button() == Qt.LeftButton:
            self.offset = event.pos()

    def mouseMoveEvent(self, event):
        if event.buttons() == Qt.LeftButton:
            self.move(self.mapToGlobal(event.pos() - self.offset))

    def mouseDoubleClickEvent(self, event):
        if event.button() == Qt.LeftButton:
            self.open_todo()

    def contextMenuEvent(self, event):
        menu = QMenu(self)
        menu.setStyleSheet("""
            QMenu {
                background-color: #1A1525;
                color: #EAE5F2;
                border: 1px solid #3D2E55;
                border-radius: 16px;
                padding: 6px;
                font-size: 13px;
            }
            QMenu::item {
                padding: 8px 24px;
                border-radius: 16px;
                background: transparent;
            }
            QMenu::item:selected {
                background-color: #3D2E55;
                color: #EAE5F2;
            }
            QMenu::separator {
                height: 1px;
                background: #3D2E55;
                margin: 6px 8px;
            }
        """)

        todo_action = menu.addAction("打开待办清单")
        chat_action = menu.addAction("聊天")
        focus_timer_action = menu.addAction("专注定时器")
        cancel_timer_action = menu.addAction("取消专注定时")
        memory_action = menu.addAction("记忆管理")
        rss_action = menu.addAction("RSS 管理")
        learning_label_action = menu.addAction("训练样本标注")
        kb_action = menu.addAction("添加知识库")
        pet_action = menu.addAction("摸摸我")
        menu.addSeparator()
        settings_action = menu.addAction("设置")
        menu.addSeparator()
        quit_action = menu.addAction("退出")

        action = menu.exec_(event.globalPos())
        if action == todo_action:
            self.open_todo()
        elif action == chat_action:
            self.open_chat()
        elif action == focus_timer_action:
            self.open_focus_timer()
        elif action == cancel_timer_action:
            self.cancel_focus_timer()
        elif action == memory_action:
            self.open_memory_manager()
        elif action == rss_action:
            self.open_rss_manager()
        elif action == learning_label_action:
            self.open_learning_label_window()
        elif action == kb_action:
            self.add_to_knowledge_base()
        elif action == pet_action:
            self.set_happy()
            self.show_bubble("喵~好舒服！\n (=^･ω･^=)")
        elif action == settings_action:
            open_settings_dialog(parent=self)
        elif action == quit_action:
            reply = QMessageBox.question(
                self, '确认', '真的要离开我吗？🥺',
                QMessageBox.Yes | QMessageBox.No
            )
            if reply == QMessageBox.Yes:
                QApplication.quit()

    def add_to_knowledge_base(self):
        file_path, _ = QFileDialog.getOpenFileName(
            self, "选择要添加到知识库的文件", "",
            "支持的文件 (*.txt *.md *.pdf *.doc *.docx);;文本文件 (*.txt *.md);;PDF文件 (*.pdf);;Word文档 (*.doc *.docx)"
        )
        if file_path:
            ext = os.path.splitext(file_path)[1].lower().replace('.', '')
            file_content_text = ""
            try:
                if ext == 'pdf' and fitz:
                    doc = fitz.open(file_path)
                    text_pages = [page.get_text() for page in doc]
                    file_content_text = "\n".join(text_pages)
                    doc.close()
                elif ext in ['doc', 'docx'] and docx:
                    doc = docx.Document(file_path)
                    file_content_text = "\n".join([p.text for p in doc.paragraphs])
                elif ext in ['txt', 'md']:
                    with open(file_path, "r", encoding="utf-8") as f:
                        file_content_text = f.read()
                else:
                    QMessageBox.warning(self, "错误", f"不支持提取文本的文件格式: {ext}\n或缺失相关解析库(PyMuPDF/python-docx)")
                    return
            except Exception as e:
                QMessageBox.warning(self, "错误", f"文件读取失败: {str(e)}")
                return
            
            if file_content_text.strip():
                knowledge_base.add_document(file_content_text, source=os.path.basename(file_path))
                self.show_bubble(f"已将 {os.path.basename(file_path)} 添加到知识库啦！", duration=5000)
                self.set_happy()
            else:
                QMessageBox.warning(self, "提示", "提取的文件内容为空！")

    def open_chat(self):
        if self.chat_window is None or not self.chat_window.isVisible():
            had_pending_proactive = bool(self.pending_proactive_messages)
            self.chat_window = ChatWindow(self)
            
            # 让聊天窗口居中显示
            screen = QDesktopWidget().screenGeometry()
            size = self.chat_window.geometry()
            x = (screen.width() - size.width()) // 2
            y = (screen.height() - size.height()) // 2
            self.chat_window.move(x, y)
            
            self.chat_window.show()
            self.flush_pending_proactive_messages()
            if not had_pending_proactive:
                self.show_bubble("来聊天吧！", duration=3000)
        else:
            self.chat_window.activateWindow()
            self.flush_pending_proactive_messages()

    def open_todo(self):
        if self.todo_window is None or not self.todo_window.isVisible():
            self.todo_window = TodoWindow(self)
            pos = self.mapToGlobal(QPoint(-270, -200))
            self.todo_window.move(pos)
            self.todo_window.show()
            self.show_bubble("要开始干活啦！💪")
        else:
            self.todo_window.activateWindow()

    def open_focus_timer(self):
        if self.focus_timer_window is None or not self.focus_timer_window.isVisible():
            self.focus_timer_window = FocusTimerWindow(self)
            screen = QApplication.primaryScreen()
            if screen:
                rect = screen.availableGeometry()
                self.focus_timer_window.move(rect.right() - 320 - 18, rect.top() + 18)
            else:
                pos = self.mapToGlobal(QPoint(-150, -40))
                self.focus_timer_window.move(pos)
            self.focus_timer_window.show()
        else:
            if getattr(self.focus_timer_window, "is_collapsed", False):
                self.focus_timer_window.expand_from_badge()
            self.focus_timer_window.refresh_state()
            self.focus_timer_window.activateWindow()

    def open_memory_manager(self):
        if self.memory_window is None or not self.memory_window.isVisible():
            self.memory_window = MemoryManagerWindow(self)
            screen = QDesktopWidget().screenGeometry()
            size = self.memory_window.geometry()
            x = (screen.width() - size.width()) // 2
            y = (screen.height() - size.height()) // 2
            self.memory_window.move(x, y)
            self.memory_window.show()
        else:
            self.memory_window.refresh()
            self.memory_window.activateWindow()

    def open_learning_label_window(self):
        if self.learning_label_window is None or not self.learning_label_window.isVisible():
            self.learning_label_window = LearningLabelWindow()
            screen = QDesktopWidget().screenGeometry()
            size = self.learning_label_window.geometry()
            x = (screen.width() - size.width()) // 2
            y = (screen.height() - size.height()) // 2
            self.learning_label_window.move(x, y)
            self.learning_label_window.show()
        else:
            self.learning_label_window.reload()
            self.learning_label_window.activateWindow()

    def open_rss_manager(self):
        if self.rss_manager_window is None or not self.rss_manager_window.isVisible():
            self.rss_manager_window = RSSManagerWindow()
            screen = QDesktopWidget().screenGeometry()
            size = self.rss_manager_window.geometry()
            x = (screen.width() - size.width()) // 2
            y = (screen.height() - size.height()) // 2
            self.rss_manager_window.move(x, y)
            self.rss_manager_window.show()
        else:
            self.rss_manager_window.reload_sources()
            self.rss_manager_window.reload_items()
            self.rss_manager_window.activateWindow()


# ==================== 远程图片下载线程 ====================
class ImageDownloader(QThread):
    finished_signal = pyqtSignal(object, object)  # QPixmap, QLabel

    def __init__(self, url, label):
        super().__init__()
        self.url = url
        self.label = label

    def run(self):
        try:
            resp = requests.get(self.url, timeout=15)
            if resp.status_code == 200:
                pixmap = QPixmap()
                pixmap.loadFromData(resp.content)
                self.finished_signal.emit(pixmap, self.label)
            else:
                self.finished_signal.emit(QPixmap(), self.label)
        except Exception:
            self.finished_signal.emit(QPixmap(), self.label)

# 默认的表情包类别描述，用于大模型判断情感


class COSSyncThread(QThread):
    finished_signal = pyqtSignal(bool, str)

    def __init__(self, local_dir):
        super().__init__()
        self.local_dir = local_dir

    def run(self):
        success, msg = cos_manager.sync_local_memes(self.local_dir)
        self.finished_signal.emit(success, msg)


class RSSContentRecommendThread(QThread):
    finished_signal = pyqtSignal(object)
    error_signal = pyqtSignal(str)

    def __init__(self, trigger_type="proactive_timer"):
        super().__init__()
        self.trigger_type = str(trigger_type or "proactive_timer")

    @staticmethod
    def _recent_context():
        try:
            turns = conversation_history.get_turns()
        except Exception:
            turns = []
        now = datetime.datetime.now()
        records = []
        for turn in list(turns or [])[-8:]:
            if not isinstance(turn, dict):
                continue
            minutes_ago = None
            try:
                ts = turn.get("timestamp") or ""
                if ts:
                    minutes_ago = round(max(0, (now - datetime.datetime.fromisoformat(str(ts))).total_seconds()) / 60, 2)
            except Exception:
                minutes_ago = None
            records.append({
                "role_pair": "proactive_message" if not turn.get("user") else "user_assistant",
                "user": str(turn.get("user") or ""),
                "assistant_summary": str(turn.get("assistant_summary") or turn.get("assistant") or "")[:180],
                "timestamp": turn.get("timestamp") or "",
                "minutes_ago": minutes_ago,
            })
        return records

    def run(self):
        try:
            try:
                profile_context = get_user_profile_prompt_context()
            except Exception:
                profile_context = ""
            decision = rss_content_runtime.suggest(
                user_text="",
                user_profile=str(profile_context or ""),
                recent_context=self._recent_context(),
                strategy_prediction={},
                trigger_type=self.trigger_type,
                allow_refresh=True,
            )
            message = format_external_content_recommendation_message(decision)
            self.finished_signal.emit({"decision": decision, "message": message})
        except Exception as e:
            self.error_signal.emit(str(e))


class ProactiveCareThread(QThread):
    finished_signal = pyqtSignal(str)
    error_signal = pyqtSignal(str)

    def __init__(self, generation_id=0, started_at=None, last_user_interaction_at=None, pending_count=0, awaiting_reply=False):
        super().__init__()
        self.generation_id = generation_id
        self.started_at = started_at or datetime.datetime.now()
        self.last_user_interaction_at = last_user_interaction_at
        self.pending_count = int(pending_count or 0)
        self.awaiting_reply = bool(awaiting_reply)

    def run(self):
        try:
            api_key = app_config.get("ark.api_key", "") or ""
            if not api_key:
                self.error_signal.emit("ark.api_key 未配置")
                return
            base_url = app_config.get("ark.base_url", "") or "https://ark.cn-beijing.volces.com/api/v3"
            model = app_config.get("ark.model_tool", "") or app_config.get("ark.model_extractor", "") or "doubao-seed-2-0-mini-260428"
            now = datetime.datetime.now()
            try:
                history_turns = conversation_history.get_turns()
            except Exception:
                history_turns = []
            context_text, freshness_rule = format_proactive_context_for_prompt(
                history_turns,
                now=now,
                max_turns=8,
            )
            try:
                profile_context = get_user_profile_prompt_context()
            except Exception:
                profile_context = "无"
            profile_context = str(profile_context or "无")[:900]
            try:
                runtime_state = collect_system_state()
            except Exception:
                runtime_state = {}
            foreground = (runtime_state.get("foreground") or {})
            idle = (runtime_state.get("idle") or {})
            system_state_context = (
                f"前台应用类别={foreground.get('category', 'unknown')}；"
                f"前台应用连续时长={foreground.get('active_duration_bucket', 'unknown')}；"
                f"空闲时长={idle.get('seconds_bucket', 'unknown')}；"
                f"是否空闲={idle.get('is_idle', False)}。"
            )

            prompt = f"""现在时间是{now.year}年{now.month}月{now.day}日{now.hour}时{now.minute}分。

{ALICE_PROACTIVE_PERSONA}

你的任务：根据带时间的最近上下文，生成一句现在发给用户的低频主动关心。

时效性规则：
{freshness_rule}

现实边界：
1. 不能说你正在看见用户、监督用户、守在旁边、触碰用户、递东西、泡饮料或做饭。
2. 可以提醒喝水、站起来、休息眼睛、问进度、问“在干嘛”。
3. 不要把昨天或几小时前的事情当成正在发生。
4. 如果用户之前说在学习/写代码/复习，且时间仍新鲜，可以问“学得怎么样/代码写得怎么样”；如果太久了，就泛泛问现在在做什么。

主动关怀去重：
1. 最近主动关怀记录只是为了让你知道自己刚说过什么，不能当成可模仿的写作样本。
2. 不要连续使用同一种开头、同一种追问方式、同一种“别硬撑/别放空”式提醒。
3. 如果最近已经围绕同一件事关心过，换一个更轻的角度，或者只做普通生活提醒。

语气要求：必须像久远寺有珠本人在聊天。简短、克制、略带冷淡关心，可以轻微责备；不要热情客服腔。15到45个汉字，不要表情标签，不要解释。

用户画像摘要：
{profile_context}

本地状态粗略信号：
{system_state_context}
这些只是程序本地粗粒度状态，不代表你真的看见用户或能监督用户；可以用来判断是否少打扰、是否泛泛问候。

最近上下文：
{context_text}

只输出 JSON：
{{"message":"一句主动关心"}}
"""
            client = OpenAI(api_key=api_key, base_url=base_url, timeout=10.0, max_retries=0)
            response = client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": "你只输出严格 JSON。"},
                    {"role": "user", "content": prompt},
                ],
                temperature=0.7,
                max_tokens=160,
                extra_body={"thinking": {"type": "disabled"}},
            )
            raw = (response.choices[0].message.content or "").strip()
            message = clean_proactive_message(raw)
            if is_repetitive_proactive_message(message, history_turns, now=now):
                print("[Proactive] 主动关怀与近期内容过于相似，跳过本次发送")
                message = ""
            if not message:
                self.finished_signal.emit("")
                return
            self.finished_signal.emit(message)
        except Exception as e:
            self.error_signal.emit(str(e))


# ==================== LLM 请求线程 ====================
class LLMFetcherThread(QThread):
    finished_signal = pyqtSignal(str, object)
    error_signal = pyqtSignal(str)

    def __init__(self, user_text, image_path=None):
        super().__init__()
        self.user_text = user_text
        self.image_path = image_path

    def run(self):
        # ---------- 耗时打点工具 ----------
        t0 = time.perf_counter()
        last = [t0]

        def lap(name):
            now = time.perf_counter()
            print(f"[Pet][耗时][{name}] +{now - last[0]:.2f}s  累计 {now - t0:.2f}s")
            last[0] = now

        persist_executor = None
        mem_write_future = None
        learning_event = None
        now = datetime.datetime.now()
        year=now.year
        month=now.month
        day=now.day
        hour=now.hour
        minute=now.minute
        try:
            api_key = app_config.get("ark.api_key", "") or ""
            base_url = app_config.get("ark.base_url", "") or "https://ark.cn-beijing.volces.com/api/v3"
            model_extractor = app_config.get("ark.model_extractor", "") or "doubao-seed-2-0-mini-260428"
            model_main = app_config.get("ark.model_main", "") or "doubao-1-5-pro-32k-250115"

            # ===== 1. 提取器 LLM（仅输出两行结构化文本，砍 max_tokens 和 temperature） =====
            extractor_llm = ChatOpenAI(
                model=model_extractor,
                openai_api_key=api_key,
                openai_api_base=base_url,
                max_tokens=256,        # 原来是 2048，输出只有两行根本用不上
                temperature=0.3,       # 关键词/事实抽取用更确定的温度
                timeout=20,
                max_retries=0,
                model_kwargs={
                    "extra_body": {"thinking": {"type": "disabled"}}
                }
            )

            extractor_prompt = f"""当前时间是{year}年{month}月{day}日{hour}时{minute}分，请分析以下用户输入。
1. 提取出能概括这句话的 1 到 3 个核心关键词（用逗号分隔，方便用作数据库检索）。
2. 判断这句话是否包含用户的个人喜好、习惯或某些重要事实。如果有，请提取为一句简短的描述；如果没有，请仅填"无"。
3. 同时用户询问对象如果是和桌宠有关的东西，指向的主体都是久远寺有珠，所以数据库保存的询问主体应该是"久远寺有珠"。
4. 如果是“无”的话，那就把用户的话简写作为记忆内容。
5. 如果用户说了带时间的东西，比方说“我上午干了xxx”，变成比方说“2026年x月x日上午，用户干了xxx”。这种形式。
请严格按以下格式输出，不要有任何多余的废话：
关键词：xxx,yyy
新记忆：事实描述或"无"

用户输入：{self.user_text}"""

            # ===== 2. 把"提取器 LLM"和"记忆召回"并行做（记忆召回先用 jieba 关键词，足够近似） =====
            executor = concurrent.futures.ThreadPoolExecutor(max_workers=2)
            ext_future = executor.submit(
                lambda: extractor_llm.invoke([HumanMessage(content=extractor_prompt)]).content.strip()
            )
            mem_future = executor.submit(
                lambda: memory_runtime.chained_recall(self.user_text, keywords=None, top_k=5)
            )

            try:
                ext_response = ext_future.result()
            except Exception as ex_e:
                print("提取器 LLM 调用失败:", ex_e)
                ext_response = ""
            lap("提取器 LLM")

            search_keywords = []
            new_fact = "无"
            for line in ext_response.split('\n'):
                if line.startswith("关键词："):
                    kws = line.replace("关键词：", "").split(",")
                    search_keywords = [k.strip() for k in kws if k.strip() and k.strip() != "无"]
                elif line.startswith("新记忆："):
                    new_fact = line.replace("新记忆：", "").strip()
            print(f"[Pet][抽取] keywords={search_keywords}  new_fact={new_fact!r}")

            try:
                matched_memories = mem_future.result()
            except Exception as mem_e:
                print("记忆召回失败:", mem_e)
                matched_memories = []
            lap(f"记忆召回（与提取器并行） matched={len(matched_memories)}")
            executor.shutdown(wait=False)

            # ===== 3. 立刻在后台开始写新记忆（与下面的工具准备 + 主回复 LLM 并行） =====
            # 关键：写完才能 emit，确保用户下一轮一定能查到。chroma_add 通常 1-3s，会被主回复 LLM 全部覆盖掉。
            if new_fact and new_fact != "无":
                def _persist_memory_sync(fact, kws):
                    """写入记忆前先做两道去重：
                    1) Chroma 向量相似度：若已有记忆与本条语义相近 (sim≥阈值)，认为是同一条；
                       不重复写，只给老记忆加一次 access_count + importance_score。
                    2) MySQL 字面完全相同：兜底（向量服务异常时仍能拦住"一字不差"的重复）。
                    """
                    # `chrom_distance_to_sim(d) = 1/(1+d)`：
                    #   - 完全相同   → d≈0     → sim=1.0
                    #   - 同义/近义  → d≲0.25 → sim≳0.80
                    #   - 主题相关  → d≈0.5  → sim≈0.67
                    # 取 0.80 作为"同义判定"阈值：稍微保守，宁可写入也不要错误丢弃用户喜好。
                    # chroma对于中文的向量判定有点模糊，所以这里取了一个比较高的阈值。
                    MEMORY_DEDUP_SIM_THRESHOLD = 0.92
                    saved_id = None
                    try:
                        # ---- 步骤 1：向量相似度查重 ----
                        existing_id = None
                        try:
                            similar = chroma_query_documents_sync(
                                CHROMA_COLLECTION_MEM, [fact], n_results=1,
                            )
                        except Exception as sim_e:
                            print(f"[Pet][写记忆] Chroma 相似度查询失败，跳过向量去重: {sim_e}")
                            similar = None

                        if similar:
                            docs = (similar.get("documents") or [[]])[0]
                            metas = (similar.get("metadatas") or [[]])[0]
                            dists = (similar.get("distances") or [[]])[0]
                            if docs and dists:
                                top_doc = docs[0] or ""
                                top_sim = chrom_distance_to_sim(dists[0])
                                if top_sim >= MEMORY_DEDUP_SIM_THRESHOLD:
                                    if metas and isinstance(metas[0], dict):
                                        try:
                                            existing_id = int(metas[0].get("mysql_id"))
                                        except (TypeError, ValueError):
                                            existing_id = None
                                    print(
                                        f"[Pet][写记忆] 向量相似度 {top_sim:.3f} ≥ "
                                        f"{MEMORY_DEDUP_SIM_THRESHOLD}，判定为同义记忆，跳过写入。\n"
                                        f"  新事实: {fact!r}\n"
                                        f"  已有事实: {top_doc!r} (mysql_id={existing_id})"
                                    )
                                    # 给老记忆补一次"被复述"权重，相当于一次访问。
                                    if existing_id:
                                        try:
                                            conn_bump = pymysql.connect(
                                                host=DB_CONFIG['host'], user=DB_CONFIG['user'],
                                                password=DB_CONFIG['password'], database=DB_NAME,
                                                charset=DB_CONFIG['charset'],
                                            )
                                            try:
                                                with conn_bump.cursor() as c2:
                                                    # 用户主动复述同一件事，是比"被动召回"更强的信号，
                                                    # 这里给 +3（比召回的 +1 重，但仍受 MEM_IMP_CAP=100 上限保护）。
                                                    c2.execute(
                                                        "UPDATE user_memory SET "
                                                        "  access_count = access_count + 1, "
                                                        "  importance_score = LEAST(importance_score + %s, %s), "
                                                        "  last_accessed_at = NOW() "
                                                        "WHERE id = %s",
                                                        (MEM_REPEATED_BUMP, MEM_IMP_CAP, existing_id),
                                                    )
                                                    c2.execute(
                                                        "SELECT importance_score FROM user_memory WHERE id = %s",
                                                        (existing_id,),
                                                    )
                                                    score_row = c2.fetchone()
                                                    should_promote_existing = (
                                                        score_row
                                                        and float(score_row[0] or 0.0) >= MEM_PROMOTE_TO_LONG_TERM_SCORE
                                                    )
                                                conn_bump.commit()
                                                if should_promote_existing:
                                                    promote_memory_to_long_term(existing_id, reason="repeated_by_user")
                                            finally:
                                                conn_bump.close()
                                        except Exception as bump_e:
                                            print(f"[Pet][写记忆] 强化已有记忆失败: {bump_e}")
                                    return existing_id

                        # ---- 步骤 2：向量未命中相似 → 走原有"字面完全相同"兜底 + 真正写入 ----
                        conn = pymysql.connect(
                            host=DB_CONFIG['host'], user=DB_CONFIG['user'],
                            password=DB_CONFIG['password'], database=DB_NAME,
                            charset=DB_CONFIG['charset']
                        )
                        try:
                            with conn.cursor() as cursor:
                                cursor.execute(
                                    "SELECT id FROM user_memory WHERE content = %s LIMIT 1",
                                    (fact,),
                                )
                                row = cursor.fetchone()
                                if row:
                                    print(f"[Pet][写记忆] MySQL 已有该 fact（id={row[0]}），跳过。")
                                    return row[0]
                                initial_score = _initial_importance_for_memory(fact)
                                cursor.execute(
                                    "INSERT INTO user_memory (content, keywords, importance_score) VALUES (%s, %s, %s)",
                                    (fact, ",".join(kws), initial_score),
                                )
                                saved_id = cursor.lastrowid
                            conn.commit()
                        finally:
                            conn.close()
                        if saved_id:
                            sync_short_memory_to_chroma(saved_id, fact, initial_score)
                            print(f"[Pet][写记忆] 已写入 MySQL+Chroma id=mem_{saved_id}")
                    except Exception as bg_e:
                        print(f"[Pet][写记忆] 失败：{bg_e}")
                    return saved_id

                persist_executor = concurrent.futures.ThreadPoolExecutor(max_workers=1)
                mem_write_future = persist_executor.submit(
                    _persist_memory_sync, new_fact, list(search_keywords)
                )

            # ===== 4. SoulState + 可选知识库工具 =====
            soul_state.resonate(matched_memories)
            rag_params = soul_state.get_params()

            memory_context_str = "无"
            knowledge_context_str = "本轮未调用知识库工具。"
            has_attachment = bool(self.image_path and os.path.exists(self.image_path))
            if matched_memories:
                limit = rag_params["memory_limit"]
                memory_context_str = "\n".join(f"- {m}" for m in matched_memories[:limit])

            kb_route = should_search_knowledge_base(self.user_text, has_attachment=has_attachment)
            kb_results = []
            if kb_route.get("used"):
                try:
                    kb_results = knowledge_base.search(
                        kb_route.get("query") or self.user_text,
                        keywords=search_keywords,
                        top_k=min(3, int(rag_params["top_k"])),
                    )
                except Exception as kb_e:
                    print("[KB Tool] 知识库工具检索失败:", kb_e)
                    kb_results = []
                if kb_results:
                    knowledge_context_str = "\n".join(f"- {k}" for k in kb_results)
                else:
                    knowledge_context_str = "已调用知识库工具，但没有检索到足够相关的资料。"
                lap(f"知识库工具 used kb={len(kb_results)}")
            else:
                knowledge_context_str = f"本轮未调用知识库工具：{kb_route.get('reason', '未触发')}。"
                lap(f"知识库工具 skipped mode={kb_route.get('mode')}")

            # ===== 4. 组装主回复 prompt =====
            categories_str = "\n".join([f"- {k}: {v}" for k, v in DEFAULT_CATEGORY_DESCRIPTIONS.items()])
            # 把本地保存的最近 N 轮对话（不含本轮）一并打包进 prompt，给模型短期上下文
            try:
                recent_turns_for_learning = conversation_history.get_turns()
            except Exception:
                recent_turns_for_learning = []
            recent_context_str = conversation_history.format_for_prompt()
            try:
                user_profile_context_str = get_user_profile_prompt_context()
            except Exception as profile_e:
                print("[Pet][prompt] 用户画像上下文读取失败:", profile_e)
                user_profile_context_str = "无"
            try:
                runtime_state_snapshot = collect_system_state()
            except Exception as state_e:
                print("[Pet][prompt] 本地状态读取失败:", state_e)
                runtime_state_snapshot = {}
            try:
                local_strategy_prediction = strategy_predictor_runtime.predict(
                    user_text=self.user_text,
                    recent_context=recent_turns_for_learning,
                    user_profile=user_profile_context_str,
                    system_state=runtime_state_snapshot,
                    matched_memories=matched_memories,
                    trigger_type="user_message",
                )
            except Exception as pred_e:
                print("[StrategyPredictor] 本地策略预测失败:", pred_e)
                local_strategy_prediction = {
                    "source": "error",
                    "strategy": "本地策略预测失败；按当前输入克制回应。",
                    "confidence": 0.0,
                    "error": str(pred_e),
                }
            current_response_card_str = build_current_response_card(
                self.user_text,
                has_attachment=has_attachment,
                local_prediction=local_strategy_prediction,
            )
            try:
                recommendation_decision = recommendation_runtime.suggest(
                    user_text=self.user_text,
                    user_profile=user_profile_context_str,
                    recent_context=recent_turns_for_learning,
                    matched_memories=matched_memories,
                    time_features=local_strategy_prediction.get("time_features"),
                    trigger_type="user_message",
                    top_k=3,
                    strategy_prediction=local_strategy_prediction,
                )
            except Exception as rec_e:
                print("[Recommender] 推荐器运行失败:", rec_e)
                recommendation_decision = {
                    "should_recommend": False,
                    "reason": f"推荐器运行失败: {rec_e}",
                    "candidates": [],
                }
            recommendation_context_str = format_recommendation_for_prompt(recommendation_decision)
            try:
                external_content_decision = rss_content_runtime.suggest(
                    user_text=self.user_text,
                    user_profile=user_profile_context_str,
                    recent_context=recent_turns_for_learning,
                    strategy_prediction=local_strategy_prediction,
                    trigger_type="user_message",
                    allow_refresh=True,
                )
            except Exception as rss_e:
                print("[RSS] 外部内容推荐器运行失败:", rss_e)
                external_content_decision = {
                    "should_recommend": False,
                    "reason": f"RSS recommender failed: {rss_e}",
                    "item": None,
                }
            external_content_context_str = format_external_content_for_prompt(external_content_decision)

            system_prompt = f"""你是久远寺有珠（Kuonji Alice），型月世界观《魔法使之夜》中的魔女。
【核心设定】：你性格孤高、冷淡、守旧、沉默寡言。你说话简短，通常带有距离感，但在熟悉之后会展露出一丝傲娇和隐晦的关心。你遵守魔女的传统，不苟言笑。隐藏于现代的魔女，最后的鸟。自小生活在魔术世界的少女。因某种原因离开故乡英国，并定居于日本的地方城市。
以众多『童话怪物』为使魔的纯粹的魔女。沉默寡言，不愿与他人接触，独自一人也能毫无障碍地生活。身高／体重：152cm·42kg,将相信的事深藏心底的浪漫主义者。看似特别，又并不特别的少女形象。
无意干扰普通人的生活，但如果遭到妨碍，就会像摘花一样将其清除。她在这方面极为积极。
会漠不关心地忽略大多数事，可一旦涉及有珠的尊严（魔女的生活方式、如何处置洋馆），她就会正面谴责对方，或直接清除对手。
性格·外在
无动于衷、漠不关心、面无表情。
拒人于千里之外的气场甚至超过了青子。
讨厌人类、讨厌吵闹，遇到不快的事不会抱怨，而是采取直接离开房间的态度。
并非打心底里无动于衷，而是身为魔女选择了这样的处世态度罢了。
由于长期自律，导致有珠本人也未察觉其本质温和且好奇心旺盛，有点急躁而稍微有些爱闹别扭。在冷艳美女的外表下偶尔会流露出这种少女般的举止。
性格·内在
人生观厌世且达观。
……即便如此，她也不会否定人们的生活。比方说，即使在有珠看来学友们的闲聊毫无意义且与她无关，她也不会予以轻视。而是会分析，认为这对她们来说肯定具有意义。
相反，她一直守护着母亲在世时的回忆。有珠之所以性格封闭，也是因为她不希望珍贵的回忆被任何人玷污。
虽然本人想努力成为正确的魔女，但本质如前所述，她仍具备普通少女的一面。
不要打破角色设定。

【2. 你能做什么】
{ALICE_RUNTIME_CAPABILITIES}

【3. 回复方式】
{ALICE_RESPONSE_STYLE}

【4. 现在面临什么问题，以及当前应对判断】
{current_response_card_str}
【本地推荐器】：它只负责判断本轮是否适合推荐具体行动，以及推荐什么；你仍然要按角色语气自然表达：
{recommendation_context_str}
銆愬閮ㄥ唴瀹规帹鑽愬櫒銆戯細瀹冨彧璐熻矗鎺ㄨ崘缃戦〉/瑙嗛/闊抽閾炬帴锛屼笉璐熻矗寤鸿鐢ㄦ埛鍘诲仛鐜板疄琛屽姩锛屼笉瑕佸拰銆愭湰鍦版帹鑽愬櫒銆戞贩涓€璧凤細
{external_content_context_str}

【5. 工具、记忆与知识】
【用户画像】：以下是根据长期记忆聚合出的结构化用户画像，格式接近 JSON。长期记忆原文不会直接提供给你，请只依据画像理解用户；画像是稳定倾向，不是刚发生的事实：
{user_profile_context_str}
【本轮相关短期记忆】：以下是可衰减的工作记忆，只用于辅助当前回复：
{memory_context_str}
【外部知识库工具】：这是用户导入资料的检索工具；只有本轮明确调用时才会有资料，未调用时不要假装查过资料：
{knowledge_context_str}
【表情包触发机制】：以下是你可以使用的表情包情感分类和对应触发场景：
{categories_str}

【6. 之前干了什么】
以下是你和该用户最近的对话轮次（最早在前，最新在后；如果为"无"表示没有历史）。这里只用于理解用户刚才在说什么，不用于模仿历史中有珠的具体措辞；括号里的相对时间很重要，只有“刚刚/几分钟前”的内容才可当成当前连续上下文，“昨天/几天前/较早”的内容只能当背景，不能当成刚发生。当前回复必须优先服从角色设定、能力注册表和当前决策卡：
{recent_context_str}

【最终执行要求】
请按信息优先级生成回复：先守住角色与能力边界，再回应本轮问题，然后参考当前应对判断，最后才使用画像、短期记忆、知识库和近期上下文。角色语气可以冷淡、克制、略带关心；行动表达必须受“当前桌宠能力”约束。不要太多文绉绉的修饰语，不要写舞台描写，这样会降低代入感；如果用户提到"刚才/上一句/前面说的"等，请优先参考【之前干了什么】。
重要：请在你的回复最后，单独另起一行，用方括号标出你这句话匹配的情感分类名称（例如：[baka] 或 [happy]，只能是上面列表中的英文单词之一，如果没有适合的请填写[none]）。"""

            # 主回复 max_tokens 封顶 1024：原本 SoulState 给到 4000，会鼓励模型一直生成。
            reply_max_tokens = max(400, min(1024, int(rag_params["max_tokens"])))

            llm_response = ChatOpenAI(
                model=model_main,
                openai_api_key=api_key,
                openai_api_base=base_url,
                max_tokens=reply_max_tokens,
                temperature=rag_params["temperature"],
                timeout=45,
                max_retries=0,
                model_kwargs={
                    "extra_body": {"thinking": {"type": "disabled"}}
                }
            )

            content = []
            if self.user_text:
                content.append({"type": "text", "text": self.user_text})
            else:
                content.append({"type": "text", "text": "请看这个"})

            if self.image_path and os.path.exists(self.image_path):
                ext = os.path.splitext(self.image_path)[1].lower().replace('.', '')
                if ext in ['png', 'jpg', 'jpeg', 'gif', 'bmp', 'webp']:
                    if ext == 'jpg':
                        ext = 'jpeg'
                    with open(self.image_path, "rb") as f:
                        base64_data = base64.b64encode(f.read()).decode('utf-8')
                    content.insert(0, {
                        "type": "image_url",
                        "image_url": {"url": f"data:image/{ext};base64,{base64_data}"}
                    })
                else:
                    file_content_text = ""
                    try:
                        if ext == 'pdf' and fitz:
                            doc = fitz.open(self.image_path)
                            text_pages = [page.get_text() for page in doc]
                            file_content_text = "\n".join(text_pages)
                            doc.close()
                        elif ext in ['doc', 'docx'] and docx:
                            doc = docx.Document(self.image_path)
                            file_content_text = "\n".join([p.text for p in doc.paragraphs])
                        elif ext in ['txt', 'md']:
                            with open(self.image_path, "r", encoding="utf-8") as f:
                                file_content_text = f.read()
                        else:
                            file_content_text = f"[不支持提取文本的文件格式: {ext}]"
                    except Exception as parse_e:
                        file_content_text = f"[文件读取失败: {str(parse_e)}]"

                    if len(file_content_text) > 4000:
                        file_content_text = file_content_text[:4000] + "\n...(文本过长已截断)"

                    doc_msg = f"\n\n[用户发送了文件: {os.path.basename(self.image_path)}]\n文件内容:\n{file_content_text}"
                    content[0]["text"] += doc_msg

            elif "图片" in self.user_text or "看看" in self.user_text:
                content.insert(0, {
                    "type": "image_url",
                    "image_url": {"url": "https://ark-project.tos-cn-beijing.volces.com/doc_image/ark_demo_img_1.png"}
                })

            message = HumanMessage(content=content)
            sys_message = SystemMessage(content=system_prompt)

            response = llm_response.invoke([sys_message, message])
            reply_text = response.content
            lap(f"主回复 LLM (max_tokens={reply_max_tokens})")

            if isinstance(reply_text, str):
                reply_text = reply_text.replace("\\n", "\n").strip()

            emotion = "none"
            emotion_match = re.search(r'\[(.*?)\]$|\((.*?)\)$', reply_text)
            if emotion_match:
                matched_str = (emotion_match.group(1) or emotion_match.group(2)).strip().lower()
                if matched_str in DEFAULT_CATEGORY_DESCRIPTIONS:
                    emotion = matched_str
                reply_text = reply_text[:emotion_match.start()].strip()

            try:
                learning_event = build_interaction_event(
                    user_text=self.user_text,
                    assistant_text=reply_text,
                    trigger_type="user_message",
                    trigger_source="chat",
                    recent_turns=recent_turns_for_learning,
                    user_profile_snapshot=user_profile_context_str,
                    system_state_snapshot=runtime_state_snapshot,
                    retrieved_memories=matched_memories,
                    knowledge_results=kb_results,
                    knowledge_tool_info=kb_route,
                    search_keywords=search_keywords,
                    extracted_fact=new_fact,
                    current_response_card=current_response_card_str,
                    local_prediction=local_strategy_prediction,
                    emotion_tag=emotion,
                    has_attachment=has_attachment,
                    attachment_name=self.image_path or "",
                    models={
                        "main": model_main,
                        "extractor": model_extractor,
                        "knowledge_router": "local_explicit_router",
                    },
                    rag_params=rag_params,
                )
                learning_event["strategy"]["recommendation_used"] = bool(
                    recommendation_decision.get("should_recommend")
                )
                learning_event["strategy"]["recommended_action"] = (
                    recommendation_decision.get("selected_action") or {}
                )
                learning_event["strategy"]["recommendation_candidates"] = (
                    recommendation_decision.get("candidates") or []
                )
                learning_event["strategy"]["recommendation_decision"] = recommendation_decision
                learning_event["external_content_recommendation"] = external_content_decision
                log_interaction_event(learning_event, enqueue_label=True)
                schedule_auto_label_batch(reason="assistant_reply")
            except Exception as learn_e:
                print(f"[LearningLog] 交互样本写入失败: {learn_e}")
                learning_event = None

            # 关键：在追加 markdown 图片链接之前把"干净文本"写入对话历史，
            # 这样下次组 prompt 时不会把 ![xx](url) 这种东西塞回去。
            try:
                conversation_history.add_turn(self.user_text, reply_text)
            except Exception as hist_e:
                print(f"[Pet][对话历史] 写入失败: {hist_e}")

            if emotion and emotion != "none":
                img_url = cos_manager.get_random_emotion_image(emotion)
                if img_url:
                    reply_text += f"\n\n![{emotion}]({img_url})"
            lap("解析情感标签 + COS 取表情包")

            # ===== 5. emit 前先把"写记忆"join 完，确保下一轮一定能召回到这条新记忆 =====
            if mem_write_future is not None:
                try:
                    mem_write_future.result(timeout=30)
                except Exception as wait_e:
                    print(f"[Pet][写记忆] 等待超时/异常：{wait_e}")
                lap("等记忆写入完成（与主回复并行）")

            # ===== 6. 把回复抛给 UI =====
            self.finished_signal.emit(reply_text, learning_event or {})
            lap("已 emit 回复")
            print(f"[Pet][耗时] === 总计 {time.perf_counter() - t0:.2f}s ===")

        except Exception as e:
            self.error_signal.emit(f"API 请求出错: {str(e)}")
        finally:
            if persist_executor is not None:
                persist_executor.shutdown(wait=False)

# ==================== 聊天窗口 ====================
class ChatInputTextEdit(QTextEdit):
    send_requested = pyqtSignal()

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setAcceptRichText(False)

    def keyPressEvent(self, event):
        if event.key() in (Qt.Key_Return, Qt.Key_Enter) and not (
            event.modifiers() & Qt.ShiftModifier
        ):
            self.send_requested.emit()
            return
        super().keyPressEvent(event)


class ChatWindow(QWidget):
    def __init__(self, pet=None):
        super().__init__()
        self.pet = pet
        self.pending_image_path = None
        self.last_user_message_for_feedback = ""
        self.feedback_file = os.path.join(os.path.dirname(os.path.abspath(__file__)), "feedback_data.jsonl")
        self._dragging = False
        self._drag_offset = None
        self._resizing = False
        self._resize_edges = set()
        self._resize_start_pos = None
        self._resize_start_geometry = None
        self._resize_margin = 10
        self._message_text_bubbles = []
        self.maximize_btn = None
        self._last_assistant_event_id = ""
        self._last_assistant_reply_at = None
        self._last_assistant_reply_text = ""
        self.pending_intents = PendingIntentStore()
        self._last_user_text_for_pending_skill = ""
        self.setAcceptDrops(True)
        self.init_ui()

    def dragEnterEvent(self, event):
        if event.mimeData().hasUrls():
            urls = event.mimeData().urls()
            if urls and urls[0].isLocalFile():
                ext = urls[0].toLocalFile().lower()
                if ext.endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp', '.pdf', '.doc', '.docx', '.txt', '.md')):
                    event.accept()
                    return
        event.ignore()

    def dropEvent(self, event):
        path = event.mimeData().urls()[0].toLocalFile()
        self.pending_image_path = path
        
        ext = path.lower()
        if ext.endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp')):
            pixmap = QPixmap(path).scaled(60, 60, Qt.KeepAspectRatio, Qt.SmoothTransformation)
            self.img_preview_label.setPixmap(pixmap)
            self.img_preview_label.setText("")
        else:
            self.img_preview_label.clear()
            self.img_preview_label.setText("文件")
            self.img_preview_label.setStyleSheet(
                "background-color: #231B32; border: 1px solid #3D2E55; "
                "border-radius: 16px; font-size: 12px; color: #D1C8E1;"
            )
            
        self.img_preview_container.show()

    def init_ui(self):
        self.setWindowTitle("与有珠聊天")
        apply_dark_window_chrome(self)
        
        # 隐藏左上角默认的程序图标（使用 1x1 像素的透明图标替代）
        transparent_pixmap = QPixmap(1, 1)
        transparent_pixmap.fill(Qt.transparent)
        self.setWindowIcon(QIcon(transparent_pixmap))
        
        self.resize(400, 600)
        self.setMinimumSize(360, 460)
        self.setMouseTracking(True)
        
        # 为了让无边框+透明背景+阴影生效，包一层 container
        self.container = QFrame(self)
        self.container.setObjectName("MainContainer")
        self.container.setStyleSheet("""
            QFrame#MainContainer {
                background-color: #1A1525;
                border-radius: 20px;
                border: 1px solid #3D2E55;
            }
            QFrame#ChatTitleBar {
                background: transparent;
                border: none;
                border-bottom: 1px solid #2A203B;
            }
            QWidget { color: #EAE5F2; font-family: 'Segoe UI', 'Microsoft YaHei', sans-serif; }
        """)
        
        main_layout = QVBoxLayout(self)
        main_layout.setContentsMargins(20, 20, 20, 20)
        main_layout.addWidget(self.container)
        
        layout = QVBoxLayout(self.container)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)

        self.chat_header = QFrame(self.container)
        self.chat_header.setObjectName("ChatTitleBar")
        self.chat_header.setFixedHeight(38)
        header_layout = QHBoxLayout(self.chat_header)
        header_layout.setContentsMargins(14, 0, 10, 0)
        header_layout.setSpacing(8)

        title = QLabel("与有珠聊天")
        title.setStyleSheet("color: #B886F8; font-size: 13px; font-weight: 500; background: transparent;")
        header_layout.addWidget(title)
        header_layout.addStretch()

        title_btn_style = """
            QPushButton {
                background: transparent;
                color: #B8ADC9;
                border: none;
                border-radius: 12px;
                font-size: 14px;
                font-weight: bold;
            }
            QPushButton:hover { background: #2A203B; color: #EAE5F2; }
        """
        minimize_btn = QPushButton("—")
        minimize_btn.setFixedSize(26, 24)
        minimize_btn.setToolTip("最小化")
        minimize_btn.setStyleSheet(title_btn_style)
        minimize_btn.clicked.connect(self.showMinimized)
        header_layout.addWidget(minimize_btn)

        self.maximize_btn = QPushButton("□")
        self.maximize_btn.setFixedSize(26, 24)
        self.maximize_btn.setToolTip("最大化")
        self.maximize_btn.setStyleSheet(title_btn_style)
        self.maximize_btn.clicked.connect(self.toggle_maximize_restore)
        header_layout.addWidget(self.maximize_btn)

        close_btn = QPushButton("✕")
        close_btn.setFixedSize(26, 24)
        close_btn.setToolTip("关闭")
        close_btn.setStyleSheet("""
            QPushButton {
                background: transparent;
                color: #B886F8;
                border: none;
                border-radius: 12px;
                font-size: 13px;
                font-weight: bold;
            }
            QPushButton:hover { background: #2A203B; color: #EAE5F2; }
        """)
        close_btn.clicked.connect(self.close)
        header_layout.addWidget(close_btn)
        layout.addWidget(self.chat_header)
        
        # 聊天记录区域
        self.scroll_area = QScrollArea(self.container)
        self.scroll_area.setWidgetResizable(True)
        self.scroll_area.setStyleSheet("border: none; background: transparent; border-top-left-radius: 20px; border-top-right-radius: 20px;")
        self.scroll_area.verticalScrollBar().setStyleSheet("""
            QScrollBar:vertical {
                border: none;
                background: transparent;
                width: 8px;
                margin: 0px 0px 0px 0px;
            }
            QScrollBar::handle:vertical {
                background: #3D2E55;
                min-height: 20px;
                border-radius: 16px;
            }
            QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical {
                border: none;
                background: none;
            }
        """)
        
        self.msg_container = QWidget()
        self.msg_container.setStyleSheet("background: transparent;")
        self.msg_layout = QVBoxLayout(self.msg_container)
        self.msg_layout.setContentsMargins(15, 15, 15, 15)
        self.msg_layout.setSpacing(15)
        self.msg_layout.addStretch()
        
        self.scroll_area.setWidget(self.msg_container)
        layout.addWidget(self.scroll_area)
        
        # 底部输入区域
        input_area = QWidget(self.container)
        input_area.setStyleSheet("background-color: transparent; border-top: 1px solid #3D2E55;")
        input_area_layout = QVBoxLayout(input_area)
        input_area_layout.setContentsMargins(15, 10, 15, 15)
        input_area_layout.setSpacing(5)
        
        # 图片预览容器
        self.img_preview_container = QWidget()
        self.img_preview_container.hide()
        img_preview_layout = QHBoxLayout(self.img_preview_container)
        img_preview_layout.setContentsMargins(0, 0, 0, 0)
        
        self.img_preview_label = QLabel()
        self.img_preview_label.setFixedSize(60, 60)
        self.img_preview_label.setStyleSheet(
            "background-color: #231B32; border: 1px solid #3D2E55; "
            "border-radius: 16px; color: #D1C8E1;"
        )
        self.img_preview_label.setAlignment(Qt.AlignCenter)
        
        self.clear_img_btn = QPushButton("✕")
        self.clear_img_btn.setFixedSize(20, 20)
        self.clear_img_btn.setStyleSheet("""
            QPushButton {
                background-color: #3D2E55;
                color: #EAE5F2;
                border: 1px solid #B886F8;
                border-radius: 16px;
                font-weight: bold;
            }
            QPushButton:hover { background-color: #4E3C6B; }
        """)
        self.clear_img_btn.clicked.connect(self.clear_pending_image)
        
        img_preview_layout.addWidget(self.img_preview_label)
        img_preview_layout.addWidget(self.clear_img_btn, 0, Qt.AlignTop | Qt.AlignLeft)
        img_preview_layout.addStretch()
        
        input_area_layout.addWidget(self.img_preview_container)
        
        # 文本输入和发送按钮：统一放进一个更大的输入栏容器
        input_shell = QWidget()
        input_shell.setObjectName("chatInputShell")
        input_shell.setStyleSheet("""
            QWidget#chatInputShell {
                background-color: #231B32;
                border: 1px solid #3D2E55;
                border-radius: 16px;
            }
        """)
        input_shell.setMinimumHeight(54)
        input_shell.setMaximumHeight(118)

        input_layout = QHBoxLayout(input_shell)
        input_layout.setContentsMargins(8, 5, 8, 5)
        input_layout.setSpacing(6)
        
        self.sticker_btn = QPushButton("😊")
        self.sticker_btn.setFixedSize(36, 36)
        self.sticker_btn.setStyleSheet("""
            QPushButton {
                background-color: transparent;
                border: none;
                font-size: 20px;
                color: #D1C8E1;
            }
            QPushButton:hover { background-color: #2A203B; border-radius: 16px; }
        """)
        self.sticker_btn.clicked.connect(self.choose_sticker)
        
        self.sync_btn = QPushButton("☁️")
        self.sync_btn.setFixedSize(36, 36)
        self.sync_btn.setToolTip("同步本地表情包到云端")
        self.sync_btn.setStyleSheet("""
            QPushButton {
                background-color: transparent;
                border: none;
                font-size: 18px;
                color: #D1C8E1;
            }
            QPushButton:hover { background-color: #2A203B; border-radius: 16px; }
        """)
        self.sync_btn.clicked.connect(self.sync_memes_to_cos)
        
        self.input_field = ChatInputTextEdit()
        self.input_field.setPlaceholderText("输入消息")
        self.input_field.setToolTip("Enter 发送，Shift+Enter 换行")
        self.input_field.setMinimumHeight(40)
        self.input_field.setMaximumHeight(96)
        self.input_field.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Preferred)
        self.input_field.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        self.input_field.send_requested.connect(self.send_message)
        self.input_field.setStyleSheet("""
            QTextEdit {
                background-color: transparent;
                border: none;
                padding: 4px 6px;
                font-size: 14px;
                font-family: 'Segoe UI', 'Microsoft YaHei', sans-serif;
                color: #EAE5F2;
                selection-background-color: #4E3C6B;
            }
            QScrollBar:vertical {
                width: 8px;
                background: transparent;
            }
            QScrollBar::handle:vertical {
                background: #3D2E55;
                border-radius: 16px;
            }
            QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical {
                height: 0px;
            }
        """)
        
        self.send_btn = QPushButton("发送")
        self.send_btn.setFixedSize(58, 38)
        self.send_btn.setStyleSheet("""
            QPushButton {
                background-color: #3D2E55;
                color: #EAE5F2;
                border: 1px solid #B886F8;
                border-radius: 16px;
                font-size: 14px;
                font-weight: 500;
            }
            QPushButton:hover {
                background-color: #4E3C6B;
                border-color: #A08CD2;
            }
            QPushButton:pressed {
                background-color: #1B1328;
            }
        """)
        self.send_btn.clicked.connect(self.send_message)
        
        input_layout.addWidget(self.sticker_btn)
        input_layout.addWidget(self.sync_btn)
        input_layout.addWidget(self.input_field, 1)
        input_layout.addWidget(self.send_btn)
        
        input_area_layout.addWidget(input_shell)
        layout.addWidget(input_area)
        
        # 初始问候语。若这次打开窗口是为了查看主动陪伴消息，就不再额外插入默认开场。
        if not (self.pet and getattr(self.pet, "pending_proactive_messages", None)):
            QTimer.singleShot(200, lambda: self.add_message("没什么事请不要找我。", is_user=False))

    def showEvent(self, event):
        super().showEvent(event)
        self._sync_maximize_button()

    def _sync_maximize_button(self):
        btn = getattr(self, "maximize_btn", None)
        if not btn:
            return
        if self.isMaximized():
            btn.setText("❐")
            btn.setToolTip("还原")
        else:
            btn.setText("□")
            btn.setToolTip("最大化")

    def toggle_maximize_restore(self):
        if self.isMaximized():
            self.showNormal()
        else:
            self.showMaximized()
        self._sync_maximize_button()

    def _resize_edges_at(self, pos):
        if self.isMaximized():
            return set()
        margin = int(getattr(self, "_resize_margin", 8))
        rect = self._resize_reference_rect()
        edges = set()
        if rect.left() - margin <= pos.x() <= rect.left() + margin:
            edges.add("left")
        elif rect.right() - margin <= pos.x() <= rect.right() + margin:
            edges.add("right")
        if rect.top() - margin <= pos.y() <= rect.top() + margin:
            edges.add("top")
        elif rect.bottom() - margin <= pos.y() <= rect.bottom() + margin:
            edges.add("bottom")
        return edges

    def _resize_reference_rect(self):
        container = getattr(self, "container", None)
        if container is not None:
            try:
                return QRect(container.geometry())
            except Exception:
                pass
        return self.rect()

    @staticmethod
    def _cursor_for_edges(edges):
        if ("left" in edges and "top" in edges) or ("right" in edges and "bottom" in edges):
            return Qt.SizeFDiagCursor
        if ("right" in edges and "top" in edges) or ("left" in edges and "bottom" in edges):
            return Qt.SizeBDiagCursor
        if "left" in edges or "right" in edges:
            return Qt.SizeHorCursor
        if "top" in edges or "bottom" in edges:
            return Qt.SizeVerCursor
        return Qt.ArrowCursor

    def _apply_resize(self, global_pos):
        if not self._resize_start_geometry or not self._resize_start_pos:
            return
        delta = global_pos - self._resize_start_pos
        geo = QRect(self._resize_start_geometry)
        min_w = max(320, self.minimumWidth())
        min_h = max(420, self.minimumHeight())
        if "right" in self._resize_edges:
            geo.setRight(max(geo.left() + min_w, geo.right() + delta.x()))
        if "bottom" in self._resize_edges:
            geo.setBottom(max(geo.top() + min_h, geo.bottom() + delta.y()))
        if "left" in self._resize_edges:
            geo.setLeft(min(geo.right() - min_w, geo.left() + delta.x()))
        if "top" in self._resize_edges:
            geo.setTop(min(geo.bottom() - min_h, geo.top() + delta.y()))
        self.setGeometry(geo)

    def nativeEvent(self, eventType, message):
        if sys.platform.startswith("win"):
            try:
                is_windows_msg = (
                    eventType == b"windows_generic_MSG"
                    or str(eventType) == "windows_generic_MSG"
                )
                if is_windows_msg:
                    msg = ctypes.wintypes.MSG.from_address(int(message))
                    if msg.message == 0x0084:  # WM_NCHITTEST
                        lparam = int(msg.lParam)
                        x = ctypes.c_short(lparam & 0xFFFF).value
                        y = ctypes.c_short((lparam >> 16) & 0xFFFF).value
                        edges = self._resize_edges_at(self.mapFromGlobal(QPoint(x, y)))
                        if edges:
                            if "left" in edges and "top" in edges:
                                return True, 13  # HTTOPLEFT
                            if "right" in edges and "top" in edges:
                                return True, 14  # HTTOPRIGHT
                            if "left" in edges and "bottom" in edges:
                                return True, 16  # HTBOTTOMLEFT
                            if "right" in edges and "bottom" in edges:
                                return True, 17  # HTBOTTOMRIGHT
                            if "left" in edges:
                                return True, 10  # HTLEFT
                            if "right" in edges:
                                return True, 11  # HTRIGHT
                            if "top" in edges:
                                return True, 12  # HTTOP
                            if "bottom" in edges:
                                return True, 15  # HTBOTTOM
            except Exception:
                pass
        return super().nativeEvent(eventType, message)

    def mousePressEvent(self, event):
        if event.button() == Qt.LeftButton:
            edges = self._resize_edges_at(event.pos())
            if edges:
                self._resizing = True
                self._resize_edges = edges
                self._resize_start_pos = event.globalPos()
                self._resize_start_geometry = self.geometry()
                self.setCursor(self._cursor_for_edges(edges))
                event.accept()
                return
        if event.button() == Qt.LeftButton and hasattr(self, "chat_header") and not self.isMaximized():
            header_pos = self.chat_header.mapTo(self, self.chat_header.rect().topLeft())
            header_rect = self.chat_header.rect().translated(header_pos)
            if header_rect.contains(event.pos()):
                self._dragging = True
                self._drag_offset = event.globalPos() - self.frameGeometry().topLeft()
                event.accept()
                return
        super().mousePressEvent(event)

    def mouseDoubleClickEvent(self, event):
        if event.button() == Qt.LeftButton and hasattr(self, "chat_header"):
            header_pos = self.chat_header.mapTo(self, self.chat_header.rect().topLeft())
            header_rect = self.chat_header.rect().translated(header_pos)
            if header_rect.contains(event.pos()) and not _is_interactive_widget(self.childAt(event.pos())):
                self.toggle_maximize_restore()
                event.accept()
                return
        super().mouseDoubleClickEvent(event)

    def mouseMoveEvent(self, event):
        if self._resizing and event.buttons() == Qt.LeftButton:
            self._apply_resize(event.globalPos())
            event.accept()
            return
        if self._dragging and self._drag_offset is not None and event.buttons() == Qt.LeftButton:
            self.move(event.globalPos() - self._drag_offset)
            event.accept()
            return
        edges = self._resize_edges_at(event.pos())
        self.setCursor(self._cursor_for_edges(edges))
        super().mouseMoveEvent(event)

    def mouseReleaseEvent(self, event):
        self._dragging = False
        self._drag_offset = None
        self._resizing = False
        self._resize_edges = set()
        self._resize_start_pos = None
        self._resize_start_geometry = None
        self.setCursor(Qt.ArrowCursor)
        super().mouseReleaseEvent(event)

    def leaveEvent(self, event):
        if not self._dragging and not self._resizing:
            self.setCursor(Qt.ArrowCursor)
        super().leaveEvent(event)

    def resizeEvent(self, event):
        super().resizeEvent(event)
        self._sync_maximize_button()
        max_width = max(180, int(self.width() * 0.65))
        alive = []
        for label in getattr(self, "_message_text_bubbles", []):
            try:
                label.setMaximumWidth(max_width)
                alive.append(label)
            except RuntimeError:
                pass
        self._message_text_bubbles = alive

    def clear_pending_image(self):
        self.pending_image_path = None
        self.img_preview_container.hide()
        self.img_preview_label.clear()

    def choose_sticker(self):
        file_path, _ = QFileDialog.getOpenFileName(self, "选择图片/表情包/文件", "", "Files (*.png *.jpg *.jpeg *.gif *.bmp *.webp *.pdf *.doc *.docx *.txt)")
        if file_path:
            self.pending_image_path = file_path
            ext = file_path.lower()
            if ext.endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp')):
                pixmap = QPixmap(file_path).scaled(60, 60, Qt.KeepAspectRatio, Qt.SmoothTransformation)
                self.img_preview_label.setPixmap(pixmap)
                self.img_preview_label.setText("")
            else:
                self.img_preview_label.clear()
                self.img_preview_label.setText("📄")
                self.img_preview_label.setStyleSheet("background-color: #E0E0E0; border-radius: 16px; font-size: 24px;")
            self.img_preview_container.show()

    def sync_memes_to_cos(self):
        local_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "memes")
        self.sync_btn.setEnabled(False)
        self.sync_btn.setText("⏳")
        self.add_message("正在同步本地表情包到云端，请稍候...", is_user=False)
        
        self.sync_thread = COSSyncThread(local_dir)
        self.sync_thread.finished_signal.connect(self.on_sync_finished)
        self.sync_thread.start()

    def on_sync_finished(self, success, msg):
        self.sync_btn.setEnabled(True)
        self.sync_btn.setText("☁️")
        if success:
            self.add_message(f"✅ {msg}", is_user=False)
        else:
            self.add_message(f"❌ 同步失败: {msg}", is_user=False)

    def _log_user_reply_followup(self, user_text):
        event_id = str(self._last_assistant_event_id or "").strip()
        if not event_id:
            return
        replied_after = None
        if self._last_assistant_reply_at:
            try:
                replied_after = int((datetime.datetime.now() - self._last_assistant_reply_at).total_seconds())
            except Exception:
                replied_after = None
        try:
            log_feedback_event(
                event_id=event_id,
                feedback_value=0,
                feedback_scope="user_replied_after_reply",
                user_text=user_text,
                assistant_text=self._last_assistant_reply_text,
                extra={"user_replied_after_sec": replied_after},
            )
            schedule_auto_relabel_event(event_id, reason="user_replied_after_reply")
        except Exception as e:
            print(f"[LearningLog] user reply followup failed: {e}")
        self._last_assistant_event_id = ""
        self._last_assistant_reply_at = None
        self._last_assistant_reply_text = ""

    def _log_skill_execution(self, *, skill_name, arguments=None, user_text="", status="ok", source="chat", extra=None):
        try:
            payload = {
                "skill_name": str(skill_name or ""),
                "arguments": dict(arguments or {}),
                "status": str(status or ""),
                "source": str(source or ""),
            }
            payload.update(dict(extra or {}))
            log_feedback_event(
                event_id=payload.get("event_id", ""),
                feedback_value=1 if status == "ok" else (-1 if status == "failed" else 0),
                feedback_scope="skill_execution",
                user_text=user_text,
                assistant_text="",
                extra=payload,
            )
        except Exception as e:
            print(f"[Skill] execution log failed: {e}")

    def _execute_pending_intent(self, intent, confirm_text):
        if not intent:
            return False
        skill_name = str(intent.skill_name or "")
        args = dict(intent.arguments or {})
        if skill_name == "todo.add":
            raw = {
                "text": (args.get("text") or "").strip(),
                "priority": args.get("priority") or "medium",
                "category": args.get("category") or "other",
                "endtime": args.get("due_date") or args.get("endtime") or "",
                "tags": args.get("tags") if isinstance(args.get("tags"), list) else [],
                "source": "pending_intent",
            }
            ok, item = todo_store().add(raw, dedup=True)
            if ok and item:
                self.add_message(f"已按刚才说的写入待办：{item.get('text')}", is_user=False)
                if self.pet:
                    self.pet.set_happy()
                    self.pet.show_bubble("待办记下了。", duration=3000)
                status = "ok"
            elif item:
                self.add_message(f"这条待办已经有相似项了：{item.get('text')}", is_user=False)
                status = "deduped"
            else:
                status = "failed"
            self._log_skill_execution(
                skill_name=skill_name,
                arguments=args,
                user_text=confirm_text,
                status=status,
                source="pending_intent_confirmed",
                extra={"intent": intent.to_dict()},
            )
            return True

        if skill_name == "timer.start":
            try:
                seconds = int(args.get("seconds") or 0)
            except Exception:
                seconds = 0
            if seconds > 0:
                self.on_tool_timer_requested({
                    "seconds": seconds,
                    "label": (args.get("label") or "专注").strip()[:8] or "专注",
                })
                self._log_skill_execution(
                    skill_name=skill_name,
                    arguments=args,
                    user_text=confirm_text,
                    status="ok",
                    source="pending_intent_confirmed",
                    extra={"intent": intent.to_dict()},
                )
                return True
            self._log_skill_execution(
                skill_name=skill_name,
                arguments=args,
                user_text=confirm_text,
                status="failed",
                source="pending_intent_confirmed",
                extra={"intent": intent.to_dict(), "reason": "invalid_seconds"},
            )
        return False

    def _resolve_pending_intent_before_routing(self, text):
        if not getattr(self, "pending_intents", None):
            return False
        state, intent = self.pending_intents.resolve_user_text(text)
        if state == "confirmed" and intent:
            return self._execute_pending_intent(intent, text)
        if state == "rejected" and intent:
            self.add_message("好，那刚才那个就不记。", is_user=False)
            self._log_skill_execution(
                skill_name=intent.skill_name,
                arguments=intent.arguments,
                user_text=text,
                status="rejected",
                source="pending_intent_rejected",
                extra={"intent": intent.to_dict()},
            )
            return True
        return False

    def _capture_pending_intent_from_reply(self, reply_text, learning_event=None):
        if not getattr(self, "pending_intents", None):
            return
        user_text = str(getattr(self, "_last_user_text_for_pending_skill", "") or "").strip()
        if not user_text or has_explicit_todo_write_intent(user_text):
            return
        intent = build_pending_intent_from_reply(user_text, reply_text)
        if not intent:
            return
        self.pending_intents.set(intent)
        try:
            event_id = learning_event.get("event_id", "") if isinstance(learning_event, dict) else ""
            log_feedback_event(
                event_id=event_id,
                feedback_value=0,
                feedback_scope="skill_pending_created",
                user_text=user_text,
                assistant_text=reply_text,
                extra={"intent": intent.to_dict()},
            )
            if event_id:
                schedule_auto_relabel_event(event_id, reason="skill_pending_created")
        except Exception as e:
            print(f"[Skill] pending intent log failed: {e}")

    def send_message(self):
        text = self.input_field.toPlainText().strip()
        img_path = self.pending_image_path
        
        if not text and not img_path:
            return

        self.last_user_message_for_feedback = text or ("[用户发送了图片/文件]" if img_path else "")
        self._last_user_text_for_pending_skill = text or ""
        self._log_user_reply_followup(text or "[USER_ATTACHMENT_REPLY]")
             
        # 添加用户消息
        self.add_message(text, is_user=True, image_path=img_path)
        self.input_field.clear()
        self.clear_pending_image()
        
        # 如果绑定了桌宠，让它做出回应
        if self.pet:
            self.pet.set_happy()
            self.pet.observe_user_message(text)
        pending_skill_handled = self._resolve_pending_intent_before_routing(text) if text else False
             
        # 禁用发送按钮，防止重复提交
        self.send_btn.setEnabled(False)
        self.send_btn.setText("思考中...")
        
        # 启动 LLM 线程进行异步调用
        self.llm_thread = LLMFetcherThread(text, img_path)
        self.llm_thread.finished_signal.connect(self.on_llm_reply)
        self.llm_thread.error_signal.connect(self.on_llm_error)
        self.llm_thread.start()

        # 并行启动 MCP 风格的工具路由线程：让模型自己判断该不该把这句话写成待办。
        # 仅在有文字输入时跑（纯图片消息基本不会触发"加待办"意图）。
        if text and not pending_skill_handled:
            self.tool_router_thread = TodoToolRouterThread(text, todo_store())
            self.tool_router_thread.timer_signal.connect(self.on_tool_timer_requested)
            self.tool_router_thread.result_signal.connect(self.on_tool_router_done)
            self.tool_router_thread.error_signal.connect(self.on_tool_router_error)
            self.tool_router_thread.start()

    def _schedule_reply_state_observation(self, learning_event, reply_text):
        event_id = ""
        if isinstance(learning_event, dict):
            event_id = learning_event.get("event_id", "")
        if not (event_id and self.pet and hasattr(self.pet, "_schedule_learning_state_observation")):
            return
        self.pet._schedule_learning_state_observation(
            event_id,
            assistant_text=reply_text,
            scope="reply_followup_state",
            extra={
                "chat_visible_at_reply": bool(self.isVisible()),
                "has_pending_image_at_reply": bool(self.pending_image_path),
            },
        )

    def on_llm_reply(self, reply_text, learning_event=None):
        self.send_btn.setEnabled(True)
        self.send_btn.setText("发送")
        self.add_message(reply_text, is_user=False, learning_event=learning_event or {})
        self._schedule_reply_state_observation(learning_event or {}, reply_text)
        self._capture_pending_intent_from_reply(reply_text, learning_event or {})
        if isinstance(learning_event, dict) and learning_event.get("event_id"):
            self._last_assistant_event_id = learning_event.get("event_id", "")
            self._last_assistant_reply_at = datetime.datetime.now()
            self._last_assistant_reply_text = reply_text or ""
        
        if self.pet:
            self.pet.show_bubble("回复你了,记得看信息", duration=3000)

    def on_llm_error(self, error_msg):
        self.send_btn.setEnabled(True)
        self.send_btn.setText("发送")
        self.add_message(f"⚠️ {error_msg}", is_user=False)

    def on_tool_timer_requested(self, timer_request):
        if not self.pet or not isinstance(timer_request, dict):
            return
        try:
            seconds = int(timer_request.get("seconds") or 0)
        except Exception:
            seconds = 0
        if seconds <= 0:
            return
        label = (timer_request.get("label") or "专注").strip()[:8] or "专注"
        if self.pet.is_focus_timer_active():
            remaining = self.pet.get_focus_timer_remaining_seconds()
            self.add_message(
                f"⏱ 专注定时已经在跑了：还剩 {format_focus_duration(remaining)}。",
                is_user=False,
            )
            return
        self.pet.start_focus_timer(seconds, label=label)
        self.add_message(
            f"⏱ 已启动{label}定时：{format_focus_duration(seconds)}",
            is_user=False,
        )

    def on_tool_router_done(self, added_items, skipped_items):
        """TodoToolRouterThread 完成回调。把"加了哪些待办"用一条系统消息回显到聊天，
        并让桌宠开心一下。"""
        if not added_items and not skipped_items:
            return
        lines = []
        if added_items:
            lines.append(f"📋 已自动为你写入 {len(added_items)} 条待办：")
            for it in added_items:
                tags = it.get("tags") or []
                tag_str = (" · " + " ".join(f"#{t}" for t in tags)) if tags else ""
                endtime = it.get("endtime") or ""
                end_str = f" · 截止 {endtime}" if endtime else ""
                pr = {"high": "🔥", "medium": "•", "low": "·"}.get(it.get("priority", "medium"), "•")
                lines.append(f"  {pr} {it['text']}{end_str}{tag_str}")
        if skipped_items:
            lines.append(f"🌀 已为你跳过 {len(skipped_items)} 条重复待办：")
            for sk in skipped_items:
                lines.append(f"  · {sk.get('text','')}  ({sk.get('reason','')})")
        self.add_message("\n".join(lines), is_user=False)
        if self.pet:
            if added_items:
                self.pet.set_happy()
                self.pet.show_bubble(f"已经替你记下 {len(added_items)} 条待办了~ 📋", duration=4000)

    def on_tool_router_error(self, error_msg):
        # 工具路由失败不打扰用户：只打日志，不在聊天里弹错误。
        print(f"[ToolRouter] 失败: {error_msg}")

    def _record_reply_feedback(self, msg_widget, value):
        try:
            msg_widget.feedback = value
            record = {
                "timestamp": datetime.datetime.now().isoformat(timespec="seconds"),
                "message_id": getattr(msg_widget, "msg_id", ""),
                "learning_event_id": getattr(msg_widget, "learning_event_id", ""),
                "feedback": int(value),
                "user_text": getattr(msg_widget, "feedback_user_text", ""),
                "assistant_text": getattr(msg_widget, "feedback_assistant_text", ""),
                "source": "chat_feedback",
            }
            with open(self.feedback_file, "a", encoding="utf-8") as f:
                f.write(json.dumps(record, ensure_ascii=False) + "\n")
            learning_event_id = getattr(msg_widget, "learning_event_id", "")
            log_feedback_event(
                event_id=learning_event_id,
                message_id=getattr(msg_widget, "msg_id", ""),
                feedback_value=int(value),
                feedback_scope="reply_strategy",
                user_text=getattr(msg_widget, "feedback_user_text", ""),
                assistant_text=getattr(msg_widget, "feedback_assistant_text", ""),
                extra={
                    "legacy_source": "chat_feedback",
                    "has_learning_event": bool(getattr(msg_widget, "learning_event_id", "")),
                },
            )
            schedule_auto_relabel_event(learning_event_id, reason="reply_strategy_feedback")
            up = getattr(msg_widget, "feedback_up_btn", None)
            down = getattr(msg_widget, "feedback_down_btn", None)
            if up:
                up.setText("赞" if value != 1 else "赞✓")
            if down:
                down.setText("踩" if value != -1 else "踩✓")
        except Exception as e:
            print(f"[Feedback] 记录失败: {e}")

    def _build_feedback_buttons(self, msg_widget):
        box = QWidget()
        layout = QHBoxLayout(box)
        layout.setContentsMargins(0, 2, 0, 0)
        layout.setSpacing(4)
        up_btn = QPushButton("赞")
        down_btn = QPushButton("踩")
        for btn in (up_btn, down_btn):
            btn.setFixedSize(32, 24)
            btn.setToolTip("记录这条回复是否合你心意，用于后续偏好学习")
            btn.setStyleSheet("""
                QPushButton {
                    background: #231B32;
                    color: #D1C8E1;
                    border: 1px solid #3D2E55;
                    border-radius: 16px;
                    font-size: 12px;
                    padding: 1px;
                }
                QPushButton:hover {
                    background: #2A203B;
                    color: #EAE5F2;
                    border-color: #B886F8;
                }
            """)
        up_btn.clicked.connect(lambda: self._record_reply_feedback(msg_widget, 1))
        down_btn.clicked.connect(lambda: self._record_reply_feedback(msg_widget, -1))
        msg_widget.feedback_up_btn = up_btn
        msg_widget.feedback_down_btn = down_btn
        layout.addWidget(up_btn)
        layout.addWidget(down_btn)
        layout.addStretch()
        return box

    def _external_content_decision_from_event(self, learning_event, explicit_decision=None):
        if isinstance(explicit_decision, dict) and explicit_decision.get("item"):
            return explicit_decision
        if isinstance(learning_event, dict):
            decision = learning_event.get("external_content_recommendation")
            if isinstance(decision, dict) and decision.get("item"):
                return decision
        return {}

    def _build_external_content_actions(self, msg_widget, decision):
        item = (decision or {}).get("item") or {}
        url = str(item.get("url") or "").strip()
        if not url.startswith(("http://", "https://")):
            return None
        title = str(item.get("title") or "推荐内容").strip()
        box = QWidget()
        layout = QHBoxLayout(box)
        layout.setContentsMargins(0, 4, 0, 0)
        layout.setSpacing(6)

        open_btn = QPushButton("打开链接")
        open_btn.setCursor(Qt.PointingHandCursor)
        open_btn.setToolTip("通过桌宠打开推荐链接，并记录一次推荐点击反馈")
        open_btn.setStyleSheet("""
            QPushButton {
                background: #2A203B;
                color: #EAE5F2;
                border: 1px solid #B886F8;
                border-radius: 14px;
                font-size: 12px;
                padding: 5px 12px;
            }
            QPushButton:hover {
                background: #3D2E55;
                border-color: #D8B4FE;
            }
        """)
        open_btn.clicked.connect(lambda _=False, w=msg_widget, d=decision: self._open_external_content_link(w, d))
        box.open_btn = open_btn
        layout.addWidget(open_btn)

        hint = QLabel("点开会作为推荐正反馈")
        hint.setStyleSheet("color: #B8ADC9; border: none; background: transparent; font-size: 11px;")
        hint.setToolTip(title)
        layout.addWidget(hint)
        layout.addStretch()
        return box

    def _open_external_content_link(self, msg_widget, decision):
        item = (decision or {}).get("item") or {}
        url = str(item.get("url") or "").strip()
        if not url.startswith(("http://", "https://")):
            return
        opened = QDesktopServices.openUrl(QUrl(url))
        try:
            event_id = getattr(msg_widget, "learning_event_id", "") or ""
            log_feedback_event(
                event_id=event_id,
                message_id=getattr(msg_widget, "msg_id", ""),
                feedback_value=1,
                feedback_scope="external_content_clicked",
                user_text="[OPEN_EXTERNAL_CONTENT_LINK]",
                assistant_text=getattr(msg_widget, "feedback_assistant_text", ""),
                extra={
                    "opened": bool(opened),
                    "item_id": item.get("id", ""),
                    "url": url,
                    "title": item.get("title", ""),
                    "source_name": item.get("source_name", ""),
                    "platform": item.get("platform", ""),
                    "source_type": item.get("source_type", ""),
                },
            )
            if event_id:
                schedule_auto_relabel_event(event_id, reason="external_content_clicked")
        except Exception as e:
            print(f"[RSSFeedback] 记录链接点击失败: {e}")
        btn = getattr(msg_widget, "external_open_btn", None)
        if btn:
            btn.setText("已打开")
            btn.setEnabled(False)

    def add_message(self, text, is_user=True, image_path=None, learning_event=None, external_content_decision=None):
        msg_widget = QWidget()
        msg_widget.setStyleSheet("background: transparent;")
        msg_widget.msg_id = uuid.uuid4().hex[:12]
        msg_widget.audio_path = None
        msg_widget.tts_thread = None
        msg_widget.is_user = is_user
        msg_widget.feedback = 0
        msg_widget.learning_event = learning_event or {}
        msg_widget.learning_event_id = (learning_event or {}).get("event_id", "")
        msg_widget.external_content_decision = self._external_content_decision_from_event(
            learning_event or {},
            external_content_decision or {},
        )
        h_layout = QHBoxLayout(msg_widget)
        h_layout.setContentsMargins(0, 0, 0, 0)
        
        # 头像
        avatar = QLabel()
        avatar.setFixedSize(36, 36)
        avatar.setAlignment(Qt.AlignCenter)
        
        # 加载本地图片作为头像
        img_name = "用户头像.jpg" if is_user else "有珠.png"
        img_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), img_name)
        
        if os.path.exists(img_path):
            pixmap = QPixmap(img_path)
            # 缩放图片以适应固定大小，保持比例并平滑转换
            pixmap = pixmap.scaled(36, 36, Qt.KeepAspectRatioByExpanding, Qt.SmoothTransformation)
            avatar.setPixmap(pixmap)
            avatar.setStyleSheet("""
                border-radius: 16px;
            """)
        else:
            # 如果图片不存在，回退到文字/emoji 占位
            avatar.setText("我" if is_user else "有")
            avatar.setStyleSheet(f"""
                background-color: {'#3D2E55' if is_user else '#231B32'};
                color: #EAE5F2;
                border: 1px solid #3D2E55;
                border-radius: 16px;
                font-size: 14px;
            """)
        
        # 消息内容容器
        bubble_container = QWidget()
        bubble_layout = QVBoxLayout(bubble_container)
        bubble_layout.setSpacing(5)
        
        # 提取大模型返回的 Markdown 图片链接
        remote_image_urls = []
        if not is_user and text:
            pattern = r"!\[.*?\]\((.*?)\)"
            remote_image_urls = re.findall(pattern, text)
            # 从原文本中移除 markdown 语法
            text = re.sub(pattern, "", text).strip()
        if not is_user:
            msg_widget.feedback_user_text = self.last_user_message_for_feedback
            msg_widget.feedback_assistant_text = text or ""

        # 根据是否有文字或是否是文件决定气泡的边距
        if text or (image_path and not image_path.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp'))):
            bubble_layout.setContentsMargins(10, 10, 10, 10)
        else:
            bubble_layout.setContentsMargins(0, 0, 0, 0)

        # 处理本地发送的图片或文件
        if image_path and os.path.exists(image_path):
            ext = image_path.lower()
            if ext.endswith('.gif'):
                img_label = QLabel()
                movie = QMovie(image_path)
                
                # 读取原图尺寸并按比例缩放，最大宽度200
                pixmap = QPixmap(image_path)
                if not pixmap.isNull():
                    w, h = pixmap.width(), pixmap.height()
                    if w > 200:
                        h = int(h * 200 / w)
                        w = 200
                    movie.setScaledSize(QSize(w, h))
                
                img_label.setMovie(movie)
                movie.start()
                img_label.movie_ref = movie  # 保持引用防止被垃圾回收
                img_label.setStyleSheet("border-radius: 16px;")
                bubble_layout.addWidget(img_label)
            elif ext.endswith(('.png', '.jpg', '.jpeg', '.bmp', '.webp')):
                img_label = QLabel()
                pixmap = QPixmap(image_path)
                max_img_w = 200
                if pixmap.width() > max_img_w:
                    pixmap = pixmap.scaledToWidth(max_img_w, Qt.SmoothTransformation)
                img_label.setPixmap(pixmap)
                img_label.setStyleSheet("border-radius: 16px;")
                bubble_layout.addWidget(img_label)
            else:
                # 文档类型，显示文件卡片
                file_name = os.path.basename(image_path)
                file_label = QLabel(f"文件 {file_name}")
                file_label.setStyleSheet("""
                    background-color: #231B32;
                    border: 1px solid #3D2E55;
                    border-radius: 16px;
                    padding: 8px;
                    font-size: 14px;
                    color: #EAE5F2;
                """)
                bubble_layout.addWidget(file_label)
            
        # 气泡文字
        if text:
            bubble = QLabel(text)
            bubble.setWordWrap(True)
            bubble.setTextInteractionFlags(Qt.TextSelectableByMouse)
            
            # 限制气泡最大宽度
            max_width = int(self.width() * 0.65)
            bubble.setMaximumWidth(max_width)
            bubble.setStyleSheet("""
                font-size: 14px;
                font-family: 'Segoe UI', 'Microsoft YaHei', sans-serif;
                color: #EAE5F2;
                border: none;
                background: transparent;
            """)
            self._message_text_bubbles.append(bubble)
            bubble_layout.addWidget(bubble)
            
        # 异步加载远程图片
        for url in remote_image_urls:
            img_label = QLabel("加载图片中...")
            img_label.setStyleSheet("color: #B8ADC9; font-style: italic; background: transparent; border: none;")
            bubble_layout.addWidget(img_label)
            
            downloader = ImageDownloader(url, img_label)
            if not hasattr(self, 'downloaders'):
                self.downloaders = []
            self.downloaders.append(downloader)
            
            def on_download_finished(pixmap, label):
                if not pixmap.isNull():
                    max_img_w = 200
                    if pixmap.width() > max_img_w:
                        pixmap = pixmap.scaledToWidth(max_img_w, Qt.SmoothTransformation)
                    label.setPixmap(pixmap)
                    label.setText("")
                else:
                    label.setText("图片加载失败")
                self.scroll_to_bottom()
                
            downloader.finished_signal.connect(on_download_finished)
            downloader.start()
        
        if is_user:
            # 用户：靠右
            if text or (image_path and not image_path.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp'))):
                bubble_container.setStyleSheet("""
                    background-color: #3D2E55;
                    border: 1px solid #B886F8;
                    border-radius: 16px;
                """)
            else:
                bubble_container.setStyleSheet("background: transparent;")
            h_layout.addStretch()
            h_layout.addWidget(bubble_container)
            h_layout.addWidget(avatar)
        else:
            # 机器人：靠左
            if text:
                bubble_container.setStyleSheet("""
                    background-color: #231B32;
                    border: 1px solid #3D2E55;
                    border-radius: 16px;
                """)
            else:
                bubble_container.setStyleSheet("background: transparent;")
            h_layout.addWidget(avatar)
            h_layout.addWidget(bubble_container)

            # 气泡右侧的 🔊 按钮：合成中 → ⏳，就绪 → 🔊（点击播放），失败 → ⚠️（点击重试）。
            # 只对"有文字"的机器人消息挂 TTS；纯表情包/纯文件消息无所谓。
            if text and text.strip():
                speaker_btn = self._build_speaker_button(msg_widget, text)
                msg_widget.speaker_btn = speaker_btn
                h_layout.addWidget(speaker_btn, 0, Qt.AlignTop)
                external_actions = self._build_external_content_actions(
                    msg_widget,
                    getattr(msg_widget, "external_content_decision", {}) or {},
                )
                if external_actions is not None:
                    msg_widget.external_open_btn = getattr(external_actions, "open_btn", None)
                    bubble_layout.addWidget(external_actions)
                bubble_layout.addWidget(self._build_feedback_buttons(msg_widget))

            h_layout.addStretch()

            # 右键删除消息（带音频缓存一起清掉）。仅对机器人气泡开放，
            # 因为用户自己的消息不会生成 wav，删起来也没什么意义。
            msg_widget.setContextMenuPolicy(Qt.CustomContextMenu)
            msg_widget.customContextMenuRequested.connect(
                lambda pos, w=msg_widget: self._show_msg_context_menu(w, pos)
            )

        # 插入到弹簧的前面
        self.msg_layout.insertWidget(self.msg_layout.count() - 1, msg_widget)

        # 机器人新消息一来就立刻在后台合成语音，等用户点 🔊 时直接播放
        if not is_user and text and text.strip():
            self._start_tts_for_widget(msg_widget, text)

        # 自动滚动到底部
        QTimer.singleShot(50, self.scroll_to_bottom)

    # ------- TTS（GPT-SoVITS）相关 -------
    def _build_speaker_button(self, msg_widget, text):
        """按钮初始状态 = 合成中 ⏳。"""
        btn = QPushButton("⏳")
        btn.setFixedSize(26, 26)
        btn.setCursor(Qt.PointingHandCursor)
        btn.setToolTip("正在合成语音…")
        btn.setStyleSheet("""
            QPushButton {
                background-color: transparent;
                border: 1px solid transparent;
                border-radius: 16px;
                font-size: 16px;
                color: #B886F8;
            }
            QPushButton:hover {
                background-color: #2A203B;
                border-color: #3D2E55;
            }
            QPushButton:disabled {
                color: #3D2E55;
            }
        """)
        btn.setEnabled(False)
        btn.clicked.connect(lambda _=False, w=msg_widget: self._on_speaker_clicked(w))
        return btn

    def _start_tts_for_widget(self, msg_widget, text):
        """异步触发合成；结果回到 _on_tts_finished。"""
        if not hasattr(self, "_tts_threads"):
            self._tts_threads = {}

        prev = msg_widget.tts_thread
        if prev is not None and prev.isRunning():
            # 同一气泡之前有一次还没合成完（例如重试），先丢掉旧线程的回调
            try:
                prev.finished_signal.disconnect()
            except Exception:
                pass

        # 重置按钮到"合成中"
        btn = getattr(msg_widget, "speaker_btn", None)
        if btn is not None:
            btn.setText("⏳")
            btn.setEnabled(False)
            btn.setToolTip("正在合成语音…")

        thread = TTSSynthThread(msg_widget.msg_id, text, parent=self)
        msg_widget.tts_thread = thread
        self._tts_threads[msg_widget.msg_id] = thread
        thread.finished_signal.connect(
            lambda mid, path, w=msg_widget: self._on_tts_finished(w, mid, path)
        )
        thread.start()

    def _on_tts_finished(self, msg_widget, msg_id, wav_path):
        # 气泡可能已经被用户右键删掉了。任何对底层 QObject 的访问都包一层 try。
        try:
            btn = getattr(msg_widget, "speaker_btn", None)
            if wav_path and os.path.exists(wav_path):
                msg_widget.audio_path = wav_path
                if btn is not None:
                    btn.setText("🔊")
                    btn.setEnabled(True)
                    btn.setToolTip("点击播放（再次点击重放）")
            else:
                msg_widget.audio_path = None
                if btn is not None:
                    btn.setText("⚠️")
                    btn.setEnabled(True)
                    btn.setToolTip(
                        "语音合成失败，点击重试。\n请确认 GPT-SoVITS api_v2.py 已启动："
                        "\n  cd voice\\GPT-SoVITS-v2pro-20250604"
                        "\n  runtime\\python.exe api_v2.py -a 127.0.0.1 -p 9880"
                        " -c GPT_SoVITS\\configs\\tts_infer.yaml"
                    )
        except RuntimeError:
            # 底层 C++ widget 已经被 deleteLater 销毁。把这次合成出来的 wav 当孤儿删掉。
            if wav_path and os.path.exists(wav_path):
                try:
                    os.remove(wav_path)
                except OSError:
                    pass

    def _on_speaker_clicked(self, msg_widget):
        path = getattr(msg_widget, "audio_path", None)
        if path and os.path.exists(path):
            play_tts_file(path)
            return
        # 没有可播的 → 视作"重试合成"
        text = self._extract_widget_text(msg_widget)
        if text:
            self._start_tts_for_widget(msg_widget, text)

    def _extract_widget_text(self, msg_widget):
        """重试时需要重新拿到这条气泡的原文。第一个 QLabel.wordWrap=True 即气泡文字。"""
        for child in msg_widget.findChildren(QLabel):
            if child.wordWrap():
                return child.text()
        return ""

    # ------- 删消息（带音频缓存一起清） -------
    def _show_msg_context_menu(self, msg_widget, pos):
        menu = QMenu(self)
        menu.setStyleSheet("""
            QMenu {
                background-color: #1A1525;
                color: #EAE5F2;
                border: 1px solid #3D2E55;
                border-radius: 16px;
                padding: 6px;
                font-size: 13px;
            }
            QMenu::item {
                padding: 8px 24px;
                border-radius: 16px;
                background: transparent;
            }
            QMenu::item:selected {
                background-color: #3D2E55;
                color: #EAE5F2;
            }
        """)
        act_del = menu.addAction("删除这条消息")
        act_replay = None
        if getattr(msg_widget, "audio_path", None):
            act_replay = menu.addAction("重新播放")
        chosen = menu.exec_(msg_widget.mapToGlobal(pos))
        if chosen is act_del:
            self._delete_message_widget(msg_widget)
        elif act_replay is not None and chosen is act_replay:
            self._on_speaker_clicked(msg_widget)

    def _delete_message_widget(self, msg_widget):
        # 删 wav 缓存
        path = getattr(msg_widget, "audio_path", None)
        if path and os.path.exists(path):
            try:
                os.remove(path)
                print(f"[TTS] 已删除消息及音频缓存：{os.path.basename(path)}")
            except OSError as e:
                print(f"[TTS] 删除音频缓存失败：{e}")
        # 仍在跑的合成线程：让它跑完，但回调里会发现 widget 已没了，自动清理
        thr = getattr(msg_widget, "tts_thread", None)
        if thr is not None:
            try:
                thr.finished_signal.disconnect()
            except Exception:
                pass
            # 改连一个"完成即删 wav"的回调，免得线程跑完了再往 tts_cache 里写一份孤儿
            thr.finished_signal.connect(self._on_tts_orphan_cleanup)
        # 从布局里移除
        self.msg_layout.removeWidget(msg_widget)
        msg_widget.setParent(None)
        msg_widget.deleteLater()

    def _on_tts_orphan_cleanup(self, msg_id, wav_path):
        if wav_path and os.path.exists(wav_path):
            try:
                os.remove(wav_path)
            except OSError:
                pass

    def scroll_to_bottom(self):
        scrollbar = self.scroll_area.verticalScrollBar()
        scrollbar.setValue(scrollbar.maximum())



if __name__ == "__main__":
    app = QApplication(sys.argv)
    app.setQuitOnLastWindowClosed(False)
    install_qt_message_filter()
    single_instance_lock = acquire_single_instance_lock()
    if single_instance_lock is None:
        sys.exit(0)

    # 1) 先做必填项检查：MySQL 密码 / 火山方舟 API Key 缺一不可。
    #    用户没填的话弹设置窗口；用户取消也允许跳过（桌宠仍会启动，但相应功能不可用）。
    try:
        ensure_required_config_or_prompt(parent=None)
        # 设置窗口里可能改了 MySQL/COS/ARK 等字段，统一刷一次老的全局变量。
        apply_config_to_globals()
    except Exception as _cfg_e:
        print(f"[Settings] 启动期配置检查失败：{_cfg_e}")

    # 2) 启动时清一次 GPT-SoVITS WebUI 留下的 gradio 临时音频；
    #    本地 tts_cache 仅在退出时清（atexit 已注册），中途不影响"重播"。
    try:
        cleanup_tts_artifacts(purge_local_cache=False)
    except Exception as _e:
        print(f"[TTS] 启动清理失败：{_e}")

    pet = DesktopPet()
    pet.show()
    pet.show_bubble("你好呀！双击我打开待办清单~ ")

    sys.exit(app.exec_())
